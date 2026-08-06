# -*- coding: utf-8 -*-
"""根据培训投影笔记整理《机器视觉培训笔记》docx。"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "机器视觉培训笔记.docx"


def set_run_font(run, size=11, bold=False, color=None, name="微软雅黑"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_format(p, space_after=6, space_before=0, line=1.15, first_line=None):
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = line
    if first_line is not None:
        pf.first_line_indent = Cm(first_line)


def add_bookmark(paragraph, name):
    """给段落加书签，便于目录跳转。"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
    tag = run._r
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(abs(hash(name)) % 100000))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), start.get(qn("w:id")))
    tag.addprevious(start)
    tag.addnext(end)


def heading(doc, text, level=1, bookmark=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size={1: 18, 2: 14, 3: 12}.get(level, 11), bold=True, name="微软雅黑")
    if bookmark:
        add_bookmark(p, bookmark)
    return p


def para(doc, text, bold=False, bullet=False, indent=0):
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
    else:
        p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=bold)
    set_paragraph_format(p, space_after=4)
    return p


def formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True, color=RGBColor(0xC0, 0x39, 0x2B))
    set_paragraph_format(p, space_after=8, space_before=4)
    return p


def note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run("注：" + text)
    set_run_font(run, size=10, color=RGBColor(0x55, 0x55, 0x55))
    set_paragraph_format(p, space_after=6)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()

    # 页面与默认样式
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)

    # ========== 封面 ==========
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("机器视觉培训笔记")
    set_run_font(run, size=28, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("3D线激光 · 相机 · 镜头 · 光源 · Lua · 四大算法 · 对位1+0")
    set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

    tip = doc.add_paragraph()
    tip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tip.add_run("（根据课堂投影笔记整理）")
    set_run_font(run, size=10, color=RGBColor(0x88, 0x88, 0x88))

    doc.add_page_break()

    # ========== 目录 ==========
    heading(doc, "目录", 1)
    toc_items = [
        "一、3D线激光机",
        "二、相机",
        "三、镜头",
        "四、光源",
        "五、Lua语法",
        "六、四大算法（灰度匹配、轮廓匹配、直线拟合、块状物）",
        "七、对位1+0",
        "附录：平面度测量原理",
    ]
    for item in toc_items:
        para(doc, item)
    doc.add_page_break()

    # ========== 一、3D线激光机 ==========
    heading(doc, "一、3D线激光机", 1, "ch1")

    heading(doc, "1.1 成像原理", 2)
    para(doc, "3D线激光成像原理：激光三角法。", bullet=True)
    para(doc, "成像时，相机、被测物、激光线构成一个三角形。", bullet=True)

    heading(doc, "1.2 常用品牌与架构差异", 2)
    para(doc, "常用品牌：基恩士（Keyence）、LMI（乐姆迈）、深视智能（SSZN）。", bullet=True)
    para(doc, "基恩士 / 深视智能：3D相机负责数据采集，控制器负责生成三维数据。", bullet=True)
    para(doc, "LMI：没有控制器；由 LMI 的 3D 相机同时完成数据采集与三维数据生成。", bullet=True)

    heading(doc, "1.3 编码器", 2)
    para(doc, "编码器是高精度的“尺子”。", bullet=True)
    para(doc, "作用：将角位移或线位移转换成电信号；与电机相连，将电机转角转换成脉冲信号，反馈给伺服驱动器实现闭环控制。", bullet=True)
    para(doc, "编码器分辨率 Re：", bullet=True)
    formula(doc, "Re = 导程 / 每转脉冲数 = 步距 / 脉冲反馈值　　单位：mm/pulse")
    para(doc, "信号类型：5V 两相差分信号（A、B 相）。A+ 与 A- 为一组差分，B+ 与 B- 为一组差分。", bullet=True)

    heading(doc, "1.4 输入模式（脉冲有效触发）", 2)
    para(doc, "一相一递增：A 相上升沿触发。", bullet=True)
    para(doc, "两相一递增：A、B 相同时上升沿触发。", bullet=True)
    para(doc, "两相两递增：A 或 B 相上升沿触发。", bullet=True)
    para(doc, "两相四递增：A 或 B 相上升沿或下降沿触发。", bullet=True)
    note(doc, "实际使用中只用两相四递增。")

    heading(doc, "1.5 选型依据", 2)
    formula(doc, "MR（量程）> 产品高度　　FoV（视野）> 产品宽度")

    heading(doc, "1.6 坐标系定义", 2)
    para(doc, "X 轴：与激光线方向平行。", bullet=True)
    para(doc, "Y 轴：物体运动方向。", bullet=True)

    heading(doc, "1.7 参数计算（重点）", 2)
    para(doc, "Rx：X 方向单像素精度（分辨率），表示激光线上测量点的水平间距。相机选定后 Rx 固定。", bullet=True)
    para(doc, "N：脉冲触发间隔（驱动设置中称“细化点数”），且 N 必须为整数。", bullet=True)
    formula(doc, "N = Rx / Re")
    para(doc, "Ry：Y 方向单像素精度（轮廓扫描间隔 / Y 间隔）。通过改 N 调节 Ry，理想情况 Ry = Rx。", bullet=True)
    formula(doc, "Ry = N × Re")
    para(doc, "若 Ry > Rx，图像压缩；若 Ry < Rx，图像拉伸。", bullet=True)
    para(doc, "H：扫描行数（驱动设置中称“批处理点数”）。", bullet=True)
    formula(doc, "H = (起始位置 − 结束位置) / Ry − 30")
    para(doc, "S：最大扫描速度；F 为帧率（驱动设置中称“采样周期”）。", bullet=True)
    formula(doc, "S = Ry × F × 90%")

    heading(doc, "提高帧率 F 的方法", 3)
    para(doc, "基恩士：压缩 Z 方向测量范围可提高 F。", bullet=True)
    para(doc, "LMI：压缩 X、Z 方向测量范围可提高 F。", bullet=True)

    heading(doc, "Rx 会变化的两种情况（间隔 Interval）", 3)
    para(doc, "OFF：实际 Rx = 选型 Rx", bullet=True)
    para(doc, "1/2：实际 Rx = 2 × 选型 Rx", bullet=True)
    para(doc, "1/4：实际 Rx = 4 × 选型 Rx", bullet=True)
    note(doc, "当帧率 F > 4kHz 时会出现上述情况。")

    heading(doc, "1.8 硬件通讯示意", 2)
    para(doc, "控制器 →（运动指令下发）→ 伺服驱动器 →（脉冲下发）→ 电机（含编码器）", bullet=True)
    para(doc, "编码器 →（编码器脉冲信号反馈）→ 3D 相机", bullet=True)
    para(doc, "伺服驱动器 →（脉冲触发）→ 3D 相机", bullet=True)

    doc.add_page_break()

    # ========== 二、相机 ==========
    heading(doc, "二、相机", 1, "ch2")

    heading(doc, "2.1 相机四大分类", 2)
    para(doc, "芯片工艺类型：", bold=True)
    para(doc, "CCD 相机（电耦合器件感光光电传感器）", bullet=True)
    para(doc, "CMOS 相机（互补性金属氧化物半导体器件）", bullet=True)
    note(doc, "CCD 均为全局快门；CMOS 既有全局也有卷帘。")
    para(doc, "传感器结构类型：线阵相机、面阵相机。", bullet=True)
    para(doc, "图像模式：黑白相机、彩色相机。", bullet=True)
    para(doc, "输出信号方式：模拟相机、数字相机。", bullet=True)

    heading(doc, "2.2 面阵相机曝光与分辨率", 2)
    para(doc, "曝光方式：全局曝光、卷帘曝光。", bullet=True)
    para(doc, "常见像素分辨率：", bold=True)
    add_table(
        doc,
        ["像素", "分辨率", "备注"],
        [
            ["30万", "640×480", ""],
            ["120万", "1280×960", ""],
            ["200万", "1600×1200", ""],
            ["500万", "2592×1944", "卷帘"],
            ["500万", "2448×2048", "全局"],
        ],
    )
    para(doc, "响应频率范围：可见光（普通）相机、红外相机、紫外相机等。", bullet=True)

    heading(doc, "2.3 线阵 / 面阵、像元与靶面", 2)
    para(doc, "线阵相机特点：像素呈行排列。", bullet=True)
    para(doc, "面阵相机特点：像素呈矩阵排列。", bullet=True)
    para(doc, "像元：组成 Sensor 的最小感光单元，为正方形。", bullet=True)
    para(doc, "像素：图像最小组成单元；像元与像素一一对应。", bullet=True)
    para(doc, "像元尺寸：单个像元实际物理尺寸；常见 2.2μm、3.45μm、2.4μm、1.85μm 等。", bullet=True)
    para(doc, "AOI 功能：可只使能部分像元，获取部分图像。", bullet=True)
    formula(doc, "β = Sensor长 / FoV长 = Sensor短 / FoV短　　→　　FoV = Sensor / β")
    para(doc, "通常情况下 Sensor 长:宽 = 4:3。", bullet=True)
    para(doc, "靶面尺寸：Sensor 对角线尺寸，由像元尺寸与分辨率共同决定；如 1\"（1英寸）。", bullet=True)
    note(doc, "在传感器中：1 Inch = 16mm；在数学中：1 Inch = 25.4mm。")
    para(doc, "Sensor 面积越大：价格越高；可用较大放大倍率得到相同 FoV；进光量更多，感光更好。", bullet=True)
    formula(doc, "像元尺寸 = Sensor长 / 分辨率长边对应像素点数（也可用短边计算）")
    para(doc, "同尺寸芯片、相同光照与参数下：像元尺寸越大 → 灵敏度越高 → 感光越好 → 图像越亮。", bullet=True)
    para(doc, "Sensor 尺寸：Sensor 长边和短边的实际物理尺寸。", bullet=True)

    heading(doc, "2.4 分辨率、帧率与曝光方式", 2)
    para(doc, "分辨率：Sensor 上的像元个数，或相机每次采集图像的像素总数。", bullet=True)
    para(doc, "帧率：相机采集频率，表示采集图像的速度。", bullet=True)
    para(doc, "面阵相机：帧率，单位 fps（帧/秒），指每秒最多能采多少帧。", bullet=True)
    note(doc, "彩色图三通道、灰度图单通道；彩色图帧率应按除以 3 计算。")
    para(doc, "线阵相机：行频，单位 Hz，指每秒能采集多少行。", bullet=True)
    para(doc, "全局曝光（帧曝光）：传感器阵列所有像元同时曝光，一般用于拍摄运动物体。", bullet=True)
    para(doc, "卷帘曝光（行曝光）：同一行像元同时曝光，不同行起始时间不同，逐行曝光；一般用于拍摄静止物体。", bullet=True)

    heading(doc, "飞拍与曝光时间", 3)
    para(doc, "飞拍（运动中拍摄）：选全局快门相机。", bullet=True)
    formula(doc, "单像素精度 = FoV长 / 分辨率长边对应像素点数")
    para(doc, "参数设置本质：单次曝光时间内物体移动量不超过一个像素尺寸。", bullet=True)
    formula(doc, "曝光时间 ≤ 精度 ÷ 速度")

    heading(doc, "2.5 接口、供电、曝光与增益", 2)
    para(doc, "常见相机传输接口：USB3.0、GigE/网口、1394a/b、CoaXPress(CXP)、CameraLink、10GigE 等。", bullet=True)
    para(doc, "常见镜头接口：C、CS、F、M42、M58 等。", bullet=True)
    para(doc, "相机供电：POE（以太网供电）、非 POE。", bullet=True)
    para(doc, "使用 POE 前提：相机支持 POE；采集卡也支持 POE。", bullet=True)
    para(doc, "曝光（Exposure Time）：单位 μs（厂家）/ ms（V2 & V3）；每个像元接受光信号的过程叫曝光，所花时间叫曝光时间（快门速度）。", bullet=True)
    para(doc, "增益（Gain）：对传感器信号放大的倍数；调增益图像变亮，噪声也会被放大。", bullet=True)

    heading(doc, "2.6 GigE 数据包与巨帧", 2)
    para(doc, "数据帧结构：MAC地址 → IP → 协议（含 GigE 协议）→ Payload → CRC 校验 → 帧间隔。", bullet=True)
    para(doc, "Packet Size：中间部分为有效净荷（payload）。", bullet=True)
    para(doc, "巨帧使能与否影响 Packet Size 有效范围：", bold=True)
    para(doc, "使能巨帧：46～9014 字节", bullet=True)
    para(doc, "不使能：46～1500 字节", bullet=True)
    para(doc, "海康相机最大 Packet Size：8164 字节。", bullet=True)
    para(doc, "Basler 相机最大 Packet Size：8192 字节。", bullet=True)

    heading(doc, "2.7 网口与相机配置", 2)
    add_table(
        doc,
        ["项目", "网口", "相机（驱动软件版本要配套）"],
        [
            ["IP地址", "设置网口 IP", "设置固定 IP，且与网口同一网段"],
            ["网口4项 / PacketSize", "巨帧 9014；接收缓存 2048；传输缓存 2048；中断节流率：极值", "PacketSize"],
        ],
    )
    para(doc, "使用 Basler 相机：网口驱动建议用 Basler 网口驱动；除网口 IP 外，其他用默认参数（网口四项默认不要改）。", bullet=True)
    para(doc, "备注：一台工控机可接多个相机；要求每个相机单独一个网段。", bullet=True)
    para(doc, "常识：连接相机前关闭防火墙，关闭电脑休眠。", bullet=True)

    doc.add_page_break()

    # ========== 三、镜头 ==========
    heading(doc, "三、镜头", 1, "ch3")

    heading(doc, "3.1 基本光学参数", 2)
    para(doc, "工作距离 WD：清晰成像条件下，镜头前端到物体表面的距离。", bullet=True)
    para(doc, "视野 FoV：视觉系统能看到的物理空间尺寸，例如 40×30mm。", bullet=True)
    para(doc, "景深 DoF：物方深度方向上能清晰成像的范围。", bullet=True)
    para(doc, "放大倍率 β = Sensor长边 / FoV长边 = Sensor短边 / FoV短边。", bullet=True)
    note(doc, "放大倍率不限于远心镜头，CCTV 镜头同样使用。")
    para(doc, "常见 CCTV 焦距 f：16mm、25mm、35mm、50mm。", bullet=True)
    formula(doc, "f_CCTV = β × WD")
    para(doc, "图像畸变：桶形畸变（网格外凸）、枕形畸变（网格内凹）。", bullet=True)

    heading(doc, "3.2 CCTV / FA 定焦镜头与远心镜头", 2)
    para(doc, "CCTV/FA 成像模型：小孔成像。", bullet=True)
    formula(doc, "1/u + 1/v = 1/f　　（u 物距，v 像距，f 焦距）")
    para(doc, "光圈：镜头内面积可变的光阑，控制通光量；用光圈系数 f/# 表示。", bullet=True)
    para(doc, "f/# 一般按 √2 倍递增；f/# 越小，光圈开口越大。", bullet=True)
    para(doc, "FA 镜头参数示例：光圈 f1.4、f2、f2.8、f4、f8、f16；C 接口；焦距 12.5mm；1:1.4 表示最大光圈；1\" 为最大支持靶面。", bullet=True)
    para(doc, "光圈作用：控制通光量、调节亮度；调焦环作用：控制像距、调节清晰度。", bullet=True)
    para(doc, "远心镜头成像模型：平行光。参数示例：靶面 1/1.8\"，工作距离 128mm，倍率 0.188。", bullet=True)

    heading(doc, "3.3 远心镜头分类", 2)
    para(doc, "物侧远心、像侧远心、双侧远心。", bullet=True)
    para(doc, "远心镜头特点：", bold=True)
    para(doc, "平行光输入，没有视角误差", bullet=True)
    para(doc, "镜头大于实际 FOV", bullet=True)
    para(doc, "同轴光输入", bullet=True)
    para(doc, "价钱较贵", bullet=True)
    para(doc, "低失真", bullet=True)
    para(doc, "工件距离固定", bullet=True)

    heading(doc, "3.4 景深影响因素与附件", 2)
    para(doc, "景深影响因素：光圈、工作距离、焦距。", bold=True)
    para(doc, "光圈越大，景深越小", bullet=True)
    para(doc, "焦距越大，景深越小", bullet=True)
    para(doc, "工作距离越大，景深越大", bullet=True)

    para(doc, "接圈的作用：", bold=True)
    para(doc, "与 CCTV 镜头配套，缩短最小对焦距离，最远对焦距离也会变化。", bullet=True)
    para(doc, "一个 5mm 接圈可将 CS 接口镜头转换为 C 接口镜头。", bullet=True)
    para(doc, "增加接圈会减小成像景深。", bullet=True)

    para(doc, "扩倍镜的作用：", bold=True)
    para(doc, "与 CCTV 配套，在不改变工作距离情况下，按倍率减小视野。", bullet=True)
    para(doc, "增加扩倍镜会减小通光量；同样条件下需要更亮光源或更长曝光。", bullet=True)

    heading(doc, "3.5 镜头分类、分辨率与转接部件", 2)
    para(doc, "按功能：定焦、变焦、定光圈。", bullet=True)
    para(doc, "按用途：FA、远心、线扫。", bullet=True)
    para(doc, "按视角：普通、广角、远摄。", bullet=True)
    para(doc, "按焦距：短焦、中焦、长焦。", bullet=True)
    para(doc, "<16mm 短焦：不利于高精度测量；>50mm 长焦：WD 大、FoV 小、景深小。", bullet=True)
    para(doc, "镜头分辨率：单位毫米内能分辨的线对数，单位 lp/mm；测量卡 ISO12233、USAF1951。", bullet=True)
    para(doc, "分辨力/解析力：相邻线对间隙差，代表镜头能否解析到的像元尺寸。", bullet=True)
    para(doc, "棱镜作用：平移光路、转折光路、分光棱镜。", bullet=True)
    para(doc, "接圈：缩小 FA 镜头 WD。", bullet=True)
    para(doc, "转接环：接口转换。", bullet=True)
    para(doc, "调焦环：用于线扫二次调焦。", bullet=True)

    heading(doc, "3.6 工业镜头选型原则（应用场景）", 2)
    para(doc, "定位：视野小（畸变影响小）、精度要求不高 → CCTV 首选；远心可用视清 WWH、WWK 系列。", bullet=True)
    para(doc, "检测：精度不高用 CCTV；精度高用远心。", bullet=True)
    para(doc, "测量：远心，如视清 DTCA 系列。", bullet=True)

    doc.add_page_break()

    # ========== 四、光源 ==========
    heading(doc, "四、光源", 1, "ch4")

    heading(doc, "4.1 光的本质与光谱", 2)
    para(doc, "光的本质：一种发散的电磁波，传播需要介质。", bullet=True)
    para(doc, "光的三原色：红、绿、蓝（RGB）。", bullet=True)
    para(doc, "白光由 RGB 合成，波长约 400～700nm。", bullet=True)
    para(doc, "380～400nm、700～780nm 为过渡带。", bullet=True)
    para(doc, "紫外光：穿透力极弱但易聚焦能量，可用于光刻；工业应用如 ITO 定位、UV 胶水完整性、隐形码检测。", bullet=True)
    para(doc, "红外光：穿透力极强，常见于人体扫描、CT；工业上可用于过滤产品表面特征干扰。", bullet=True)
    add_table(
        doc,
        ["波段", "范围"],
        [
            ["紫外光", "10nm～380nm"],
            ["过渡带", "380～400nm"],
            ["蓝光 B", "400～500nm"],
            ["绿光 G", "500～600nm"],
            ["红光 R", "600～700nm"],
            ["过渡带", "700～780nm"],
            ["红外光", "780nm～1mm"],
            ["可见光 / 白光", "400～700nm"],
        ],
    )

    heading(doc, "4.2 相似色、互补色与光源作用", 2)
    para(doc, "相似色：构成上有相同部分但不完全相同，色相环上相近。", bullet=True)
    para(doc, "互补色：构成上无相同部分，色相环上相对。", bullet=True)
    para(doc, "应用：合理运用相似色与互补色可控制图像对比度。", bullet=True)
    para(doc, "结论：同色打白，异色打黑；白色背景反射所有颜色光，黑色背景吸收所有颜色光。", bullet=True)
    para(doc, "光源作用：给相机补光。", bold=True)
    para(doc, "工业检测角度：①提高背景与特征对比度；②提升系统稳定性，随时为工件提供亮度；③提升抗干扰性。", bullet=True)
    para(doc, "作为参考：提高卷帘的抗震性。", bullet=True)
    para(doc, "视觉常用照明：白炽灯、荧光灯、卤素灯、LED 灯。", bullet=True)
    para(doc, "光源选择：一般选 LED（寿命长、颜色多、成本低、响应快、环保）。", bullet=True)

    heading(doc, "4.3 好图像特征与打光方式", 2)
    para(doc, "一副好图像的特征：对比度高；图像稳定且均匀性好；精度高；颜色真实；图像清晰。", bullet=True)
    para(doc, "偏振片作用：过滤大部分方向的光，使单一方向光通过；不改变光的传递方式。", bullet=True)
    para(doc, "光源角度分界线：45°。", bullet=True)
    para(doc, "高角度光源：特征暗、背景亮 —— 明场。", bullet=True)
    para(doc, "正面打光：可打出明暗场，常规检测。", bullet=True)
    para(doc, "背面打光：常用于高精度测量，特点非黑即白。", bullet=True)

    heading(doc, "4.4 明场、暗场与典型应用", 2)
    para(doc, "明场：背景打白、特征打黑（背景灰度高、特征灰度低）。", bullet=True)
    para(doc, "暗场：背景打黑、特征打白（背景灰度低、特征灰度高）。", bullet=True)
    para(doc, "同轴光：反射光线与镜头平行，被测物相当于镜子，表面凹凸会很明显。", bullet=True)
    para(doc, "平行同轴光：检测超光滑表面外观缺陷，如光盘划痕、充电器表面划痕等。", bullet=True)
    note(doc, "平行同轴光需要搭配远心镜头使用。")
    para(doc, "无影光源：①检测反光、不平整表面；②检测表面粗糙且均匀的弧面。", bullet=True)
    para(doc, "偏振光：①消除表面光斑干扰，利于特征提取；②消除薄膜表面反光；③消除高亮反光，颜色更真实。", bullet=True)

    heading(doc, "4.5 常见光源命名", 2)
    add_table(
        doc,
        ["光源类型", "缩写", "页码参考"],
        [
            ["环形光源", "RS", "P1-P5"],
            ["条形光源", "BS", "P6-P15"],
            ["条形组合光源", "ROS", "P16-17"],
            ["同轴光源", "COS", "P18-P25"],
            ["同轴平行光源", "COPS", "P28-P29"],
            ["平面同轴光源", "FCS", "P30-P31"],
            ["面光源", "FS", "P32-P54"],
            ["圆顶光源", "DS", "P55-P56"],
            ["无影光源", "FQS", "P57-P66"],
            ["点光源", "SPL", "P67-P70"],
            ["线性光源", "LPG", "P73-P90"],
            ["三色光源", "RGB", "P91-P92"],
            ["AOI光源", "SAI", "P93-P94"],
            ["紫外光源", "UV", "P95-P96"],
            ["偏光光源", "POR", "P97-P98"],
            ["红外光源", "IR", "—"],
            ["结构化光源 / 防水光源 / 工业冷光源 / 定制光源", "—", "—"],
        ],
    )

    doc.add_page_break()

    # ========== 五、Lua ==========
    heading(doc, "五、Lua语法", 1, "ch5")

    heading(doc, "5.1 注释", 2)
    para(doc, "单行注释：--", bullet=True)
    para(doc, "多行注释：--[[ 注释语句 --]]", bullet=True)

    heading(doc, "5.2 标识符（变量名 / 函数名等）", 2)
    para(doc, "以字母或下划线 _ 开头，后跟字母、下划线或数字。", bullet=True)
    para(doc, "避免使用下划线后接大写字母（如 _A、_Ade），不符合推荐命名。", bullet=True)
    para(doc, "不允许在标识符中使用 @、$、% 等特殊字符。", bullet=True)
    para(doc, "Lua 区分大小写：Runoob 与 runoob 是不同标识符。", bullet=True)
    para(doc, "正确示例：mohd、zara、abc、move_name、a_123、myname50、_temp、j、a23b9、retVal。", bullet=True)

    heading(doc, "5.3 保留关键词", 2)
    para(doc, "保留字不能作为常量、变量或其他用户自定义标识符：")
    para(doc, "and, break, do, else, elseif, end, false, for, function, if, in, local, nil, not, or, repeat, return, then, true, until, while, goto")

    heading(doc, "5.4 全局变量（标准 Lua）", 2)
    para(doc, "变量默认是全局的。", bullet=True)
    para(doc, "变量无需先声明，也无需指定数据类型；第一次赋值时自动创建。", bullet=True)
    para(doc, "访问未初始化的全局变量不会报错，结果为 nil。", bullet=True)

    heading(doc, "5.5 全局变量（MVStudio）", 2)
    para(doc, "变量非默认全局；可定义全局变量或局部变量。", bullet=True)
    para(doc, "使用前需在工具栏的全局/局部变量中先定义。", bullet=True)
    para(doc, "定义时可输入初始值，或在流程初始化中赋值。", bullet=True)

    heading(doc, "5.6 作用域（MVStudio）", 2)
    para(doc, "全局变量：应用于整个工程，任一子流程都可链接。", bullet=True)
    para(doc, "局部变量：只能用于创建时鼠标所在的子流程，其他子流程不能链接。", bullet=True)

    doc.add_page_break()

    # ========== 六、四大算法 ==========
    heading(doc, "六、四大算法（灰度匹配、轮廓匹配、直线拟合、块状物）", 1, "ch6")

    heading(doc, "6.1 通用概念（灰度匹配 & 轮廓匹配）", 2)
    para(doc, "CT（cycle time）：周期时间，加工一个工件所需耗时。", bullet=True)
    para(doc, "UPH（output per hour）：每小时产量。", bullet=True)
    para(doc, "ROI（region of interest）：感兴趣区域。", bullet=True)
    para(doc, "注册图像：也称模板图像；以注册图像中搜索对象的姿态为起始角度。", bullet=True)
    para(doc, "搜索范围/ROI：搜索对象的活动范围，可按实际活动范围优化 CT。", bullet=True)
    para(doc, "相似度：又称匹配分数，范围 0～100，越接近 100 越相似；一般默认 70。", bullet=True)
    para(doc, "高精度：又称亚像素；亚像素精度高于像素精度，只能在软件层面实现。", bullet=True)
    para(doc, "0 度定义：以注册图像 Mark 点姿态为 0 度；顺时针为负，逆时针为正。", bullet=True)
    para(doc, "角度范围略小于真实来料角度时：图形搜索任务不一定 NG，但搜索对象 ROI 与真实对象会出现不平行。", bullet=True)

    heading(doc, "6.2 灰度匹配", 2)
    para(doc, "搜索级别：又称金字塔层级，目的是提高 CT。", bullet=True)
    para(doc, "灰度匹配计算相似度方法：归一化互相关（NCC）。", bullet=True)
    para(doc, "灰度匹配基于图像灰度值特征运算。", bullet=True)
    para(doc, "灰度图单通道，彩色图三通道。", bullet=True)
    para(doc, "像素坐标系原点：左上角。", bullet=True)
    para(doc, "图像本质：带坐标的数字矩阵。", bullet=True)
    para(doc, "判定条件：相当于客户要求；不符合时流程出现红色 ×，任务执行失败。", bullet=True)
    para(doc, "图形/形状搜索 NG 原因：匹配失败、判定条件不满足、补正源任务执行失败。", bullet=True)

    heading(doc, "6.3 轮廓匹配", 2)
    para(doc, "轮廓求相似度方法：向量内积 / 点乘。", bullet=True)
    para(doc, "梯度阈值：逐像素求梯度，与阈值比较；≥阈值为有效梯度点，<阈值为无效梯度点。", bullet=True)
    para(doc, "轮廓匹配基于图像的梯度特征；梯度是一个向量。", bullet=True)
    para(doc, "最小片段尺寸：组成轮廓的像素个数若小于设定值，则该轮廓无效 —— 可去噪声。", bullet=True)
    para(doc, "比例范围：来料有尺寸缩放时使用；无缩放设 1～1；有缩放按现场设置，且上下限之和等于 2（如 0.98～1.02）。", bullet=True)
    para(doc, "优化 CT 建议调：搜索范围 ROI、角度范围、搜索级别。", bullet=True)
    note(doc, "相似度和梯度阈值也能影响 CT，但不建议调。")

    heading(doc, "6.4 直线拟合 & 块状物（补正源）", 2)
    para(doc, "引入补正源的目的：实际应用中工件在图像中位置会偏移，ROI 需跟随移动，否则可能检不到特征。通过定位基准（补正源）让 ROI 随基准联动。", bullet=True)
    para(doc, "工作原理：", bold=True)
    para(doc, "找到补正前后的基准点与基准角度。", bullet=True)
    para(doc, "根据两点、两角关系计算仿射变换矩阵。", bullet=True)
    para(doc, "原 ROI 位置信息经仿射变换得到补正后位置，再生成新 ROI。", bullet=True)
    para(doc, "使用补正源前提：相机必须固定，保证所有图像在同一坐标系。", bullet=True)
    para(doc, "求线与线距离时：线的顺序有关；从“线1”中心点作垂线。", bullet=True)

    doc.add_page_break()

    # ========== 七、对位1+0 ==========
    heading(doc, "七、对位1+0", 1, "ch7")

    heading(doc, "7.1 三个位置", 2)
    para(doc, "取料位、拍照位、放料位（组装位）。", bullet=True)

    heading(doc, "7.2 基准位获取（反向抓取）", 2)
    para(doc, "将手机壳手动放到料盘。", bullet=True)
    para(doc, "机械手移动到基准放料位后，反向抓取手机壳到拍照位。", bullet=True)
    para(doc, "相机采集并注册成基准位。", bullet=True)
    formula(doc, "实时放料位 = 基准放料位 + 偏移量 (XYR)")

    heading(doc, "7.3 两个坐标系与标定作用", 2)
    para(doc, "相机：像素坐标系。", bullet=True)
    para(doc, "机械手：物理坐标系 / 世界坐标系。", bullet=True)
    para(doc, "标定作用：①计算单像素精度（mm/pixel）；②通过仿射变换计算像素与世界坐标系关系（至少 4 个点）。", bullet=True)
    para(doc, "仿射变换示例：", bold=True)
    formula(doc, "x4' = a11·x4 + a12·y4 + a13")
    formula(doc, "y4' = a21·x4 + a22·y4 + a23")
    para(doc, "常见设定：移动标定 2×2，旋转标定 3（点）。", bullet=True)

    heading(doc, "7.4 调机流程（逆向调机）", 2)
    para(doc, "首先 PLC 有三个点位：取料位、拍照位、组装位/放料位。", bullet=True)
    para(doc, "调机时逆向调机：", bold=True)
    para(doc, "PLC 在取料位 (x1,y1,z1,θ1) 抓取手机壳后；", bullet=True)
    para(doc, "抓着手机壳到组装位 (x3,y3,z3,θ3)，慢慢调试，人眼/测试工具辅助，直到组装/放置满足要求；", bullet=True)
    para(doc, "此时不能放下手机壳；放下再抓，相对位置会变化；", bullet=True)
    para(doc, "在组装位抬起 Z 轴，角度不变，只平移 XY 到拍照位 (x2,y2,z2,θ2)；", bullet=True)
    para(doc, "到拍照位后仍只平移 XY、角度不变，尽量移到视野中央，调节 Z 满足工作距离，再调相机对焦；", bullet=True)
    para(doc, "相机拍照注册成基准图像（采集基准位）。", bullet=True)
    para(doc, "调机完成、三个点位设置后要保持不变。", bullet=True)

    heading(doc, "7.5 实时纠偏逻辑", 2)
    para(doc, "PLC 抓取实时工件到拍照位后，算法计算出偏差 (Δu1, Δv1, Δθ1)。", bullet=True)
    para(doc, "调整时必须先旋转纠正角度偏差，再平移；否则会反复调整且难到位。", bullet=True)
    para(doc, "旋转时因 PLC 与手机壳相对位置每次不同，即使旋转相同角度，新偏差 (Δu2, Δv2, Δθ2) 也不同。", bullet=True)
    para(doc, "纠偏放料时：将 (Δu1, Δv1, Δθ1) 与 (Δu2, Δv2, Δθ2) 加在一起用于组装。", bullet=True)
    para(doc, "旋转带来的新坐标偏差，只有知道旋转轴坐标才能计算出来。", bullet=True)
    note(doc, "所有有角度的纠偏，都要做旋转中心标定，也就是要将旋转轴的坐标标定出来。")

    heading(doc, "7.6 自动标定流程（PLC ↔ CCD）", 2)
    para(doc, "1）PLC 到拍照位 → CCD：开始标定信号 CBS, 0, 0, 0", bullet=True)
    para(doc, "2）CCD → PLC：偏移量 (X, Y, R) —— 物理坐标（注意加延时）", bullet=True)
    para(doc, "3）PLC 根据偏移量走位，到位后 → CCD：到位信号/拍照信号 CAE, 0, 0, 0", bullet=True)
    para(doc, "4）CCD 拍照 + 处理：", bullet=True)
    para(doc, "NG：标定失败", indent=1)
    para(doc, "OK：记录当前点位像素坐标，判断是否走完预设点位（例如预设 9 个点）", indent=1)
    para(doc, "未走完 → 回到第 2 步；走完预设点位 → 进入第 5 步", indent=1)
    para(doc, "5）标定计算，得到仿射变换矩阵", bullet=True)

    heading(doc, "7.7 移动标定", 2)
    para(doc, "移动模式：2×2；移动间隔：5mm。", bullet=True)
    para(doc, "四点呈方形分布，中心为原点；走位路径绕中心矩形边界运动。", bullet=True)
    para(doc, "绝对坐标（以 1、3 连线中点为原点）：", bold=True)
    add_table(
        doc,
        ["点位", "X", "Y", "R"],
        [
            ["1", "-2.5", "-2.5", "0"],
            ["2", "2.5", "-2.5", "0"],
            ["3", "2.5", "2.5", "0"],
            ["4", "-2.5", "2.5", "0"],
            ["原点", "0", "0", "0"],
        ],
    )
    para(doc, "相对坐标（以上一点为原点）：", bold=True)
    add_table(
        doc,
        ["点位", "X", "Y", "R"],
        [
            ["1", "-2.5", "-2.5", "0"],
            ["2", "5", "0", "0"],
            ["3", "0", "5", "0"],
            ["4", "-5", "0", "0"],
            ["回原点", "2.5", "-2.5", "0"],
        ],
    )

    heading(doc, "7.8 旋转标定", 2)
    para(doc, "旋转模式：3；旋转间隔：5°。", bullet=True)
    para(doc, "三点沿弧分布：点 2 为 0°，点 1 / 点 3 分别为正负角度。", bullet=True)
    para(doc, "绝对坐标（以点 2 为 0°）：", bold=True)
    add_table(
        doc,
        ["点位", "X", "Y", "R(°)"],
        [
            ["1", "0", "0", "5"],
            ["2", "0", "0", "0"],
            ["3", "0", "0", "-5"],
            ["原点", "0", "0", "0"],
        ],
    )
    para(doc, "相对坐标（以上一点为 0°）：", bold=True)
    add_table(
        doc,
        ["点位", "X", "Y", "R(°)"],
        [
            ["1", "0", "0", "5"],
            ["2", "0", "0", "-5"],
            ["3", "0", "0", "-5"],
            ["回原点", "0", "0", "5"],
        ],
    )

    heading(doc, "7.9 InnoVision 信号类型", 2)
    para(doc, "1：代表自动运行", bullet=True)
    para(doc, "2：代表对位计算（只有计算综合纠偏时才会用上，即两端来料都有偏差）", bullet=True)
    para(doc, "3：代表自动标定", bullet=True)

    heading(doc, "7.10 左右手系", 2)
    para(doc, "方法一：手心朝自己或俯看，大拇指指向 X 正方向，其余四指指向 Y 正方向，看哪只手符合。", bullet=True)
    para(doc, "方法二：大拇指指向 Z 正方向（朝自己），四指表示从 X 正方向转到 Y 正方向，看哪只手符合。", bullet=True)
    para(doc, "右手系示意：X 向右、Y 向上。", bullet=True)
    para(doc, "左手系示意：X 向右、Y 向下。", bullet=True)

    heading(doc, "7.11 旋转正方向与视觉/运动方向", 2)
    para(doc, "运动旋转正方向（机械手/PLC）：俯看，给正角度值，观察顺时针还是逆时针，以此定义正方向。", bullet=True)
    para(doc, "视觉与运动旋转方向规则：上同下反。", bold=True)
    para(doc, "上相机：视觉旋转方向与运动相同。", bullet=True)
    para(doc, "下相机：视觉旋转方向与运动相反。", bullet=True)

    heading(doc, "7.12 标定结果与防撞参数", 2)
    para(doc, "各点停留时间：防抖动。", bullet=True)
    para(doc, "标定结果查看：主要看 X、Y、Θ 整体误差；标定精度应在 1 个像素以内。", bullet=True)
    para(doc, "单位换算参考：1 丝 = 1 条 = 10μm = 0.01mm。", bullet=True)
    para(doc, "对位数据范围 / 判定条件：给每次对位设范围，超限则输出报警日志，防止撞机。", bullet=True)
    para(doc, "补偿设定 / 位置补偿：补偿机构机械偏差带来的输出数据误差。", bullet=True)
    para(doc, "贴合数据范围 / 综合判定条件：对位数据范围 + 补偿设定后的综合限制范围，同样用于防撞机。", bullet=True)

    heading(doc, "7.13 放置——料盘有偏差", 2)
    para(doc, "机械手抓取固定工件，移动到基准放料位，再下发信息做放料纠偏。", bullet=True)
    para(doc, "示意：CCD 下视工件；按偏置坐标（例如 -1, -1, 0）进行偏移纠正。", bullet=True)

    doc.add_page_break()

    # ========== 附录 ==========
    heading(doc, "附录：平面度测量原理", 1, "app")
    para(doc, "对测量点进行平面拟合，得到实时拟合平面。", bullet=True)
    para(doc, "dMax：距离拟合平面最大（正）距离的点。", bullet=True)
    para(doc, "dMin：距离拟合平面最小（负）距离的点。", bullet=True)
    formula(doc, "平面度 = dMax − dMin")
    note(doc, "因 dMin 为相对平面的负值，相减相当于最大正偏差与最大负偏差绝对值之和。常用于 3D 线激光测高/测平面度场景，并与直线/平面拟合算法相关。")

    footer = doc.add_paragraph()
    set_paragraph_format(footer, space_before=18)
    run = footer.add_run("整理来源：华汉伟业（HANSWELL）培训课堂投影笔记")
    set_run_font(run, size=9, color=RGBColor(0x88, 0x88, 0x88))

    doc.save(OUT)
    print(f"Wrote: {OUT}")


if __name__ == "__main__":
    build()
