# -*- coding: utf-8 -*-
"""向 demo提问.docx 追加架构/工程向答辩题与答案。"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC = Path(r"c:\Users\Admin\Desktop\dist\月度答辩\demo提问.docx")
BACKUP = SRC.with_name("demo提问.bak.docx")
OUT = SRC.with_name("demo提问-含补充.docx")
OUT_DOCS = Path(r"d:\Qt\project\opencv\docs") / "demo提问-含补充.docx"


def set_run_font(run, name="微软雅黑", size=11, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, bold=True)
    return p


def add_qa(doc, ask, answer):
    p1 = doc.add_paragraph()
    r1 = p1.add_run("追问：")
    set_run_font(r1, bold=True)
    r2 = p1.add_run("\u00a0" + ask)
    set_run_font(r2, bold=False)

    p2 = doc.add_paragraph()
    r3 = p2.add_run("答：")
    set_run_font(r3, bold=True)
    r4 = p2.add_run("\u00a0" + answer)
    set_run_font(r4, bold=False)

    doc.add_paragraph()


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=14, bold=True)
    run.font.color.rgb = RGBColor(0x1A, 0x52, 0x7A)
    doc.add_paragraph()


QA_ITEMS = [
    (
        "21. 为什么拆 ImageProcessor 和 Widget？",
        "Processor 能不能自己在参数变化时直接 reprocess？",
        "拆开是为了职责分离：ImageProcessor 只管「原图 + 块链 + 结果」，不碰 ROI 图元、会话、撤销；"
        "Widget 管 UI/ROI/会话。块参数变化只 emit requestReprocess，不能由 Processor 自己 reprocess——"
        "否则可能用到过期的 m_rois。约定永远是 Widget：getAllRoiInfo → setRois → reprocess。",
    ),
    (
        "22. 为什么块只发 requestReprocess，不直接调 reprocess？",
        "信号和直接调用差在哪？",
        "差在「谁先同步 ROI」。requestReprocess 是请求 Widget 走统一入口；若块或 Processor 直接 reprocess，"
        "拖着 ROI 改参数时可能仍用旧 mask。现场可指：所有加删块、使能、调参最终都汇到 onApplyProcessing。",
    ),
    (
        "23. BaseBlock 为什么编译进 blocksdk，而不是每个插件各带一份？",
        "主程序和插件怎么认同一个基类？",
        "主程序只编排 BaseBlock*，插件 DLL 必须和主程序链接同一份 blocksdk，才能保证 vtable/信号槽 ABI 一致。"
        "若各插件各自编译一份 BaseBlock，跨模块传指针会炸。代价：SDK 改接口后插件要一起重编。",
    ),
    (
        "24. 走一遍：DLL → 左侧列表 → 拖入 → 第一次出图",
        "热加载之后第一帧结果是怎么来的？",
        "PluginManager 扫 plugins/，QPluginLoader 按 IID 取 IBlockPlugin → 注册 id/order/工厂；"
        "左侧按 order 刷列表。拖拽 MIME 带稳定 id → createBlockByName → setHost → addBlock；"
        "若已有原图则 emit requestReprocess → Widget 同步 ROI 再跑链 → processingFinished 刷画布与耗时。",
    ),
    (
        "25. IBlockHost 是干什么的？Otsu 为何读原图？",
        "为什么不用链上当前中间图做 Otsu？",
        "IBlockHost 让插件拿「原图 + 当前 ROI」而不依赖 Widget 类型（解耦）。"
        "Otsu 经 hostOriginalImage() 取原图，有 ROI 则裁外接矩形算 T，写回 SpinBox 再触发重算。"
        "若用中间结果（已滤波/已伪彩）统计，阈值会跟上游绑死，换序后参数语义混乱；绑定原图更稳定、可预期。",
    ),
    (
        "26. 为什么整图算法再 mask 贴回，而不是只裁 ROI 算？",
        "这样有什么代价？",
        "契约简单：每个块输入输出同尺寸全图；ROI 外保持「本块输入」像素；重叠区用并集 mask，无谁优先。"
        "滤波/形态学也需要邻域，硬裁补边更麻烦。代价：大图时 ROI 虽小也要跑整图；可承认后续可对部分算子做 ROI 裁剪优化。",
    ),
    (
        "27. alignProcessedToSrc 解决什么问题？",
        "边缘检出灰度后接伪彩/ROI，会不会通道对不上？",
        "算子可能改通道（BGR↔Gray）或偶发尺寸不一致。贴回前 align：尺寸不对就 resize，通道不一致就 Gray↔BGR，"
        "再 copyTo(mask)。保证「局部处理」契约不因算子输出类型变化而破——和演示题 9 同一套机制。",
    ),
    (
        "28. 空 ROI、画了又删光、重叠边缘，mask 分别怎样？",
        "旋转矩形边缘那圈像素听谁的？",
        "空列表：RoiProcess 不做限制 ≈ 全图。删光后同样空列表。多 ROI：各 shape makeMask 后 bitwise_or 并集；"
        "mask=255 用处理后像素。旋转矩形/椭圆边缘是栅格化后的离散像素，半像素级属几何离散，没有抗锯齿优先级。",
    ),
    (
        "29. 块的所有权在谁？removeBlock 为什么只断开？",
        "谁负责 delete？",
        "所有权在 Widget（面板上的 QWidget 块）。Processor 只持有裸指针编排序列；removeBlock 从链里拿掉并 requestReprocess，"
        "真正销毁用 block->deleteLater()，避免信号处理中途删对象。",
    ),
    (
        "30. 热加载能不能卸载/热替换同一个 DLL？",
        "SDK 改了接口旧插件会怎样？",
        "当前是 load 进进程并注册工厂，没有完整 unload/热替换路径（Windows 下 DLL 已映射也难替）。"
        "现场说法：能「添加」新插件，不能热换已加载同名 DLL，需重启。SDK/BaseBlock 一变，旧插件 ABI 不兼容，要整套重编。",
    ),
    (
        "31. 扫 plugins 为何跳过 blocksdk / opencv / Qt*？",
        "误加载会怎样？",
        "这些是运行时依赖库不是 IBlockPlugin。硬加载会失败或污染列表；跳过只尝试算法插件 DLL，"
        "再用 IID 校验，非插件会 warn 并 unload。",
    ),
    (
        "32. 插件英文翻译怎么跟主程序一起切？",
        "换语言会不会改 JSON 里的块名？",
        "每个插件可带 *_en.qm；PluginManager::setEnglishUi 对已加载 translator install/remove。"
        "列表与会话 JSON 用稳定 id（= blockName()，不随 tr 变）；界面文案走 tr()。所以中英切换不破坏链恢复。",
    ),
    (
        "33. 会话何时读写？换缩略图 undo 清不清？",
        "绝对路径 key 的局限？",
        "换图：先把当前链+ROI 写入 m_sessions[绝对路径]，再按新路径 restore；落盘 sessions/app_sessions.json。"
        "换图/恢复会话时 undo 栈按实现清空或重建（避免跨图撤销串味）。文件挪路径 → key 对不上，会话丢——主动承认局限。",
    ),
    (
        "34. 撤销为何没有 Redo？恢复时为何禁 push？",
        "文档写结构撤销，拧旋钮也能撤？",
        "实现是单栈快照（最多 MAX_UNDO=40），pop 恢复，无反向 redo 栈。restore 时 m_undoRestoring=true，"
        "内部 clear/addBlock/改 ROI 不再 push，否则撤销动作本身进栈造成污染。"
        "结构变更（加删块/换序/ROI）与 paramsAboutToChange（拧参前）都会压整份会话快照；ROI 拖动中不逐帧压，"
        "只在 geometryAboutToChange 压一次。可说：产品上「结构+参数可撤」，拖动防抖与撤销粒度分开。",
    ),
    (
        "35. 大图卡顿除了 UI 线程还有哪些开销？",
        "耗时标签量的是什么？能取消吗？",
        "reprocess 在 UI 线程同步跑完整链；另有每块 QPixmap↔Mat、整图算子+mask 贴回。"
        "label 显示的是 ImageProcessor::reprocess 墙钟毫秒（块循环累计），不是缩放/刷新单独耗时。"
        "当前无进度条、无取消令牌。后续：工作线程 + 取消 + 结果回写到主线程。",
    ),
    (
        "36. 若改成异步处理，要注意什么？",
        "ROI 拖动和结果回写怎么防竞态？",
        "要点：①任务带世代号/令牌，过期结果丢弃；②ROI 快照在提交任务时拷贝，避免算到一半 mask 被改；"
        "③只有最新任务可 setResult+刷画布；④撤销/换图要取消未完成任务。答出设计即可，不必现场改代码。",
    ),
    (
        "37. GLCM 为何是「分析节点」还要进链？",
        "放在中间会不会影响后面块？",
        "process() 原样返回 input，只在面板算特征；仍参与调度、可读 ROI、可随会话保存。"
        "后面块吃的仍是上一段图像像素，视觉上几乎不变。证明链上允许非变换节点（分析/检测），"
        "和变换块同一套编排。",
    ),
    (
        "38. 复制块 vs 导出整条链 JSON，差别？",
        "粘贴进的是图还是配置？",
        "复制块：剪贴板是单块 name+saveParams JSON，粘贴 createBlock+loadParams 插到指定块后；图像不进剪贴板。"
        "导出链：整条块序列 JSON，给别人导入；缺插件则 create 失败跳过。会话落盘同构：路径 → 链 JSON + ROI。",
    ),
    (
        "39. 浅/深主题怎么做的？有没有写死颜色？",
        "换主题要不要改每个控件？",
        "StyleLoader 读 QSS（浅/深），qApp->setStyleSheet 全局套。控件用 objectName/QSS 选择器，"
        "避免业务代码写死色值；主题状态可记设置。现场切主题证明「样式与逻辑分离」。",
    ),
    (
        "40. 日志记什么？对答辩排障有什么用？",
        "日志和会话文件在哪？",
        "AppLogger 写 exe 旁 logs/；插件加载成功/失败、Otsu 阈值、各块耗时、会话读写等会打点。"
        "sessions/app_sessions.json 存跨次启动的链。演示卡顿或导入失败时可打开日志对老师说「可追溯」。",
    ),
    (
        "41. 主程序是不是空壳？加算法要改主程序吗？",
        "扩展性体现在哪？",
        "主程序负责编排：画布、ROI、会话、撤销、插件管理；算法在 plugins/*.dll。"
        "新算法实现 IBlockPlugin+BaseBlock，丢进 plugins（或设置里热添加），左侧自动出现；"
        "业务代码不写死「二值化/滤波」分支，只认稳定 id → 工厂 → BaseBlock*。",
    ),
    (
        "42. 拖拽换块顺序时，面板顺序和执行顺序如何一致？",
        "只改了 UI 没改链会怎样？",
        "列表拖拽结束会同步 Processor 里 m_blocks 顺序（与面板子控件顺序对齐），再统一重算。"
        "若只改 UI 不改引擎顺序，画面语义会和面板不一致——实现上两者绑在一起，换序立刻 reprocess 验证。",
    ),
    (
        "43. 对比模式和清空链，结果缓冲怎样处理？",
        "清空后保存/对比会不会仍是旧图？",
        "对比只切 m_showOriginal，在 original/result 两套图间显示，不重算。"
        "清空链：remove+deleteLater，resetResultToOriginal，画布打回原图，耗时归零，避免 m_result 残留。"
        "无块时对比也等价原图。",
    ),
    (
        "44. 保存结果图含不含 ROI 框？",
        "ROI 是画在图上的吗？",
        "保存的是 Processor 结果像素；ROI 是场景图元 + 蒙版逻辑，默认不烧进保存图。"
        "（若另做「导出叠加标注」需单独说明；当前按纯结果图答。）",
    ),
    (
        "45. 你认为当前最大短板和改进优先级？",
        "诚实题，怎么答加分？",
        "短板：①UI 线程同步整链，大图/GLCM/长链会卡；②整图+mask，小 ROI 也付全图成本；"
        "③插件只热加不热换；④会话绑绝对路径。优先级：工作线程+取消 → ROI 裁剪优化热点算子 → "
        "会话改相对路径/内容哈希。老师爱听「知道边界 + 有改进顺序」。",
    ),
]

EXTRA_CHAIN = [
    "补演示 5 连（架构向，可接在原连环炮后）：",
    "加块 →「谁触发重算？为何不直接 reprocess？」",
    "Otsu →「Host 原图，不是中间图」",
    "小 ROI 大图 →「整图+mask 的代价」",
    "设置添加 DLL →「工厂注册；不能热替换」",
    "拧参再 Ctrl+Z →「整份会话快照 + 恢复禁 push」",
]

EXTRA_TABLE = [
    ["统一入口", "块只 requestReprocess；Widget 先 ROI 再算"],
    ["IBlockHost", "插件取原图/ROI，不依赖 Widget 类型"],
    ["整图+mask", "契约简单；小 ROI 也有全图成本"],
    ["分析块", "GLCM 直通像素，只出统计"],
    ["撤销", "会话快照栈；恢复期 m_undoRestoring"],
    ["热加载", "可添加注册；不承诺卸载热换"],
    ["短板", "UI 线程同步；承认并给改进序"],
]


def main():
    # 优先读原文件；若原文件被占用导致读失败则用备份
    try:
        raw = SRC.read_bytes()
    except Exception:
        raw = BACKUP.read_bytes()
    BACKUP.write_bytes(raw)

    from io import BytesIO
    doc = Document(BytesIO(raw))

    add_section_heading(doc, "补充提问（架构 / 工程 / 边界）——建议与前 20 题一起背")

    for title, ask, answer in QA_ITEMS:
        add_title(doc, title)
        add_qa(doc, ask, answer)

    add_section_heading(doc, "补充「连环炮」与速记")
    for line in EXTRA_CHAIN:
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, bold=line.startswith("补演示"))

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("补充速记卡")
    set_run_font(run, bold=True)

    # 复用文档里已有表格的样式名（部分 docx 没有 "Table Grid"）
    style_name = None
    if doc.tables:
        try:
            style_name = doc.tables[0].style.name
        except Exception:
            style_name = None
    table = doc.add_table(rows=1 + len(EXTRA_TABLE), cols=2)
    if style_name:
        try:
            table.style = style_name
        except Exception:
            pass
    hdr = table.rows[0].cells
    hdr[0].text = "关键词"
    hdr[1].text = "一句话"
    for i, (k, v) in enumerate(EXTRA_TABLE):
        table.rows[i + 1].cells[0].text = k
        table.rows[i + 1].cells[1].text = v

    for path in (OUT, OUT_DOCS):
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
        print(f"Wrote: {path}")

    try:
        doc.save(str(SRC))
        print(f"Also updated original: {SRC}")
    except PermissionError:
        print(f"Original locked, please close Word and replace with: {OUT.name}")


if __name__ == "__main__":
    main()
