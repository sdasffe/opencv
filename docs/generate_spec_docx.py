# -*- coding: utf-8 -*-
"""生成《图像处理工具》需求与设计说明（docx，可点击目录）
对照当前完整实现：插件化（blocksdk + plugins/*.dll）、多 ROI、GLCM、
撤销、会话落盘、浅/深主题、热添加插件、AppLogger 等。
"""

from pathlib import Path
import math
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
FIG.mkdir(exist_ok=True)
OUT = ROOT / (
    "".join(chr(c) for c in [
        0x56FE, 0x50CF, 0x5904, 0x7406, 0x5DE5, 0x5177, 0x2D,
        0x9700, 0x6C42, 0x4E0E, 0x8BBE, 0x8BA1, 0x8BF4, 0x660E,
    ]) + ".docx"
)


# ========================= 绘图工具 =========================

def font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for p in candidates:
        if Path(p).exists():
            try:
                return ImageFont.truetype(p, size)
            except Exception:
                pass
    return ImageFont.load_default()


def rounded_rect(draw, xy, r, fill, outline=None, width=2):
    draw.rounded_rectangle(xy, radius=r, fill=fill, outline=outline, width=width)


def center_text(draw, box, text, fnt, fill=(30, 30, 30)):
    x0, y0, x1, y1 = box
    bbox = draw.textbbox((0, 0), text, font=fnt)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((x0 + x1 - tw) / 2, (y0 + y1 - th) / 2), text, font=fnt, fill=fill)


def draw_arrow(draw, start, end, color=(60, 60, 60), width=2):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    angle = math.atan2(y2 - y1, x2 - x1)
    size = 10
    left = (x2 - size * math.cos(angle - 0.4), y2 - size * math.sin(angle - 0.4))
    right = (x2 - size * math.cos(angle + 0.4), y2 - size * math.sin(angle + 0.4))
    draw.polygon([end, left, right], fill=color)


def make_block_diagram():
    """总体模块框图：主程序 + blocksdk + 插件 + 数据流"""
    w, h = 1380, 980
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small, tiny = font(22, True), font(14, True), font(12), font(11)
    d.text((28, 14), "总体模块框图（插件化架构 · 数据流与接口）", font=title_f, fill=(33, 33, 33))

    def box(x, y, ww, hh, title, lines, fill, border):
        rounded_rect(d, (x, y, x + ww, y + hh), 10, fill, border, 2)
        d.text((x + 12, y + 10), title, font=box_f, fill=border)
        for i, line in enumerate(lines):
            d.text((x + 12, y + 36 + i * 18), line, font=small, fill=(45, 45, 45))

    box(40, 50, 520, 170, "主界面模块 Widget（实现 IBlockHost）", [
        "菜单：文件/编辑/设置/帮助/关于",
        "打开图·文件夹·画布·多 ROI·撤销·会话",
        "设置：中/英 · 浅/深主题 · 添加插件",
        "接口：createBlockByName / getAllRoiInfo",
        "       onApplyProcessing / pushUndoSnapshot",
    ], "#E3F2FD", "#1565C0")

    box(590, 50, 340, 170, "算法列表 MyListWidget", [
        "左侧工具箱（插件动态填充）",
        "拖出：MIME text = 块 id",
        "稳定 id（UserRole）",
        "中英切换不丢匹配",
        "rebuildAlgoList ← PluginManager",
    ], "#E8EAF6", "#3949AB")

    box(960, 50, 380, 170, "配置 / 样式 / i18n / 日志", [
        "AppConfig 常量（缩放/MIME/色）",
        "StyleLoader：theme_light/dark.qss",
        "主程序 opencv_en.qm + 插件 *_en.qm",
        "AppLogger → exe/logs/",
        "QSettings 持久化语言/主题",
    ], "#F3E5F5", "#7B1FA2")

    box(40, 250, 420, 160, "处理调度 ImageProcessor", [
        "持有：原图 / 结果 / 链 / 多 ROI",
        "add/remove/moveBlock · setRois",
        "reprocess 按序 process",
        "信号：requestReprocess / finished(ms)",
    ], "#E8F5E9", "#2E7D32")

    box(490, 250, 420, 160, "blocksdk · PluginManager", [
        "IBlockPlugin / IBlockHost 接口",
        "扫描 plugins/*.dll 注册工厂",
        "createBlock(id) · loadPluginFile",
        "插件自带翻译 attachTranslation",
    ], "#E8EAF6", "#283593")

    box(940, 250, 400, 160, "算法链 BaseBlock 族（插件内）", [
        "抽象：process / blockName / 参数 JSON",
        "信号：paramsChanged / remove / copy…",
        "六插件：二值化/形态学/滤波/",
        "        灰度/伪彩/GLCM",
    ], "#FFF3E0", "#EF6C00")

    box(40, 445, 280, 145, "ROI 模块", [
        "RoiInfo + JSON",
        "Resizable*Item 多选",
        "RoiProcess 并集蒙版",
    ], "#FFF8E1", "#F9A825")

    box(340, 445, 280, 145, "会话 ImageSession", [
        "按图片路径存链+ROI",
        "内存 QHash + 落盘",
        "sessions/app_sessions.json",
    ], "#E0F2F1", "#00695C")

    box(640, 445, 340, 145, "plugins/block_*.dll", [
        "各自含 Block + *Algorithm",
        "实现 IBlockPlugin 导出",
        "热添加：设置→添加插件",
    ], "#E0F7FA", "#00838F")

    box(1000, 445, 340, 145, "工具 utils/", [
        "ImageConverter / TimeMeasurer",
        "RoiProcess / AppLogger",
        "链入 blocksdk 供插件链接",
    ], "#FCE4EC", "#C2185B")

    draw_arrow(d, (300, 220), (250, 250), "#1565C0", 3)
    d.text((260, 222), "setOriginal / setRois / reprocess", font=tiny, fill=(21, 101, 192))
    draw_arrow(d, (700, 220), (700, 250), "#283593", 3)
    d.text((710, 222), "ids / createBlock", font=tiny, fill=(40, 53, 147))
    draw_arrow(d, (700, 410), (900, 330), "#EF6C00", 3)
    d.text((780, 360), "工厂创建块", font=tiny, fill=(230, 81, 0))
    draw_arrow(d, (1140, 410), (1140, 445), "#00838F", 2)

    d.text((40, 620), "主数据流：原图 → [启用 Block1.process] → … → [BlockN] → 结果图 → 画布",
          font=small, fill=(40, 40, 40))
    d.text((40, 650), "控制流：参数/换序/ROI(60ms防抖) → requestReprocess → 同步 ROI → reprocess → finished",
          font=small, fill=(40, 40, 40))
    d.text((40, 680), "插件流：启动扫描 plugins/ → PluginManager 注册 → 左侧列表；设置可热加载单个 DLL",
          font=small, fill=(40, 40, 40))
    d.text((40, 710), "会话流：换图/退出 → ImageSession(chain+rois) → app_sessions.json；撤销栈≤40",
          font=small, fill=(40, 40, 40))
    d.text((40, 750), "目录：core/ · blocks/ · blocksdk/ · plugins/ · roi/ · utils/ · config/ · styles/ · i18n/",
          font=tiny, fill=(90, 90, 90))
    d.text((40, 780), "工程：opencv.pro = blocksdk → plugins → opencv_app（SUBDIRS ordered）",
          font=tiny, fill=(90, 90, 90))

    path = FIG / "20_module_block.png"
    img.save(path)
    return path


def make_flowchart():
    """主处理流程（加长版，分 A～G 阶段）"""
    w, h = 1320, 1960
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small, tiny = font(22, True), font(13, True), font(12), font(11)
    d.text((28, 10), "主流程图：启动 → 打开图像 → 会话/ROI → 建链 → 重算 → 显示/落盘",
           font=title_f, fill=(33, 33, 33))

    def node(x, y, ww, hh, text, fill, border, diamond=False):
        if diamond:
            cx0, cy0 = x + ww / 2, y + hh / 2
            pts = [(cx0, y), (x + ww, cy0), (cx0, y + hh), (x, cy0)]
            d.polygon(pts, fill=fill, outline=border)
        else:
            rounded_rect(d, (x, y, x + ww, y + hh), 8, fill, border, 2)
        lines = text.split("\n")
        if len(lines) == 1:
            center_text(d, (x, y, x + ww, y + hh), text, box_f)
        else:
            total = len(lines) * 18
            cy = y + (hh - total) / 2
            for i, line in enumerate(lines):
                bbox = d.textbbox((0, 0), line, font=box_f)
                tw = bbox[2] - bbox[0]
                d.text((x + (ww - tw) / 2, cy + i * 18), line, font=box_f, fill=(30, 30, 30))

    def varrow(x, y1, y2, color=(80, 80, 80)):
        draw_arrow(d, (x, y1), (x, y2), color, 2)

    def phase(x, y, text, color):
        d.text((x, y), text, font=box_f, fill=color)

    L, nw, nh, nd = 280, 420, 46, 56
    mid = L + nw // 2
    gap = 12
    bypass_x = 110

    y = 48
    phase(40, y, "【A 启动】", "#283593")
    y += 26
    a_steps = [
        "创建 QApplication / Fusion 样式",
        "设置大图内存上限 + AppLogger::init",
        "扫描 plugins/*.dll → PluginManager 注册",
        "加载浅/深主题 QSS",
        "创建并显示 Widget 主窗口",
        "rebuildAlgoList 填充左侧算法列表",
    ]
    a_ys = []
    for t in a_steps:
        node(L, y, nw, nh, t, "#E8EAF6", "#283593")
        a_ys.append(y)
        y += nh + gap
    for i in range(len(a_ys) - 1):
        varrow(mid, a_ys[i] + nh, a_ys[i + 1])

    y = a_ys[-1] + nh + gap + 6
    phase(40, y, "【B 打开图像】", "#1565C0")
    y += 26
    node(L, y, nw, nd, "打开单图 还是 文件夹？", "#E3F2FD", "#1565C0", True)
    y_open = y
    y = y_open + nd + gap + 6

    w2, h2 = 200, 72
    x1, x2 = L - 10, L + nw - w2 + 10
    node(x1, y, w2, h2, "单图\n选文件→loadImageFromPath", "#E3F2FD", "#1565C0")
    node(x2, y, w2, h2, "文件夹\n缩略图条→点选切换", "#E3F2FD", "#1565C0")
    y_branch = y
    d.line([(mid, y_open + nd), (x1 + w2 // 2, y_open + nd), (x1 + w2 // 2, y_branch)],
           fill=(21, 101, 192), width=2)
    draw_arrow(d, (x1 + w2 // 2, y_branch - 1), (x1 + w2 // 2, y_branch), (21, 101, 192), 2)
    d.line([(mid, y_open + nd), (x2 + w2 // 2, y_open + nd), (x2 + w2 // 2, y_branch)],
           fill=(21, 101, 192), width=2)
    draw_arrow(d, (x2 + w2 // 2, y_branch - 1), (x2 + w2 // 2, y_branch), (21, 101, 192), 2)
    d.text((x1 + 70, y_open + nd + 2), "单图", font=tiny, fill=(21, 101, 192))
    d.text((x2 + 55, y_open + nd + 2), "文件夹", font=tiny, fill=(21, 101, 192))

    y = y_branch + h2 + gap
    node(L, y, nw, nh, "setOriginalImage + 画布显示原图", "#E3F2FD", "#1565C0")
    y_setimg = y
    d.line([(x1 + w2 // 2, y_branch + h2), (x1 + w2 // 2, y_setimg - 6), (mid, y_setimg - 6)],
           fill=(21, 101, 192), width=2)
    d.line([(x2 + w2 // 2, y_branch + h2), (x2 + w2 // 2, y_setimg - 6), (mid, y_setimg - 6)],
           fill=(21, 101, 192), width=2)
    draw_arrow(d, (mid, y_setimg - 6), (mid, y_setimg), (21, 101, 192), 2)

    y = y_setimg + nh + gap + 6
    phase(40, y, "【C 会话恢复】", "#00695C")
    y += 26
    node(L, y, nw, nd, "该图片路径已有会话？", "#E0F2F1", "#00695C", True)
    y_sess = y
    y = y_sess + nd + gap + 6
    node(x1, y, w2, h2, "是：restoreSession\n重建链 + 多 ROI", "#E0F2F1", "#00695C")
    node(x2, y, w2, h2, "否：空链\n等待用户建链/ROI", "#E0F2F1", "#00695C")
    y_sess_b = y
    d.line([(mid, y_sess + nd), (x1 + w2 // 2, y_sess + nd), (x1 + w2 // 2, y_sess_b)],
           fill=(0, 105, 92), width=2)
    draw_arrow(d, (x1 + w2 // 2, y_sess_b - 1), (x1 + w2 // 2, y_sess_b), (0, 105, 92), 2)
    d.line([(mid, y_sess + nd), (x2 + w2 // 2, y_sess + nd), (x2 + w2 // 2, y_sess_b)],
           fill=(0, 105, 92), width=2)
    draw_arrow(d, (x2 + w2 // 2, y_sess_b - 1), (x2 + w2 // 2, y_sess_b), (0, 105, 92), 2)
    d.text((x1 + 80, y_sess + nd + 2), "是", font=tiny, fill=(0, 105, 92))
    d.text((x2 + 80, y_sess + nd + 2), "否", font=tiny, fill=(0, 105, 92))

    y = y_sess_b + h2 + gap
    node(L, y, nw, nh, "汇合：进入 ROI / 建链（有链时稍后重算）", "#E0F2F1", "#00695C")
    y_sess_join = y
    d.line([(x1 + w2 // 2, y_sess_b + h2), (x1 + w2 // 2, y_sess_join - 6), (mid, y_sess_join - 6)],
           fill=(0, 105, 92), width=2)
    d.line([(x2 + w2 // 2, y_sess_b + h2), (x2 + w2 // 2, y_sess_join - 6), (mid, y_sess_join - 6)],
           fill=(0, 105, 92), width=2)
    draw_arrow(d, (mid, y_sess_join - 6), (mid, y_sess_join), (0, 105, 92), 2)

    y = y_sess_join + nh + gap + 6
    phase(40, y, "【D ROI】", "#F9A825")
    y += 26
    node(L, y, nw, nd, "是否添加 / 调整 ROI？", "#FFF8E1", "#F9A825", True)
    y_roi = y
    y = y_roi + nd + gap
    node(L, y, nw, nh, "添加：矩形 / 椭圆 / 旋转矩形（可多个）", "#FFF8E1", "#F9A825")
    y_roi_add = y
    y = y_roi_add + nh + gap
    node(L, y, nw, nh, "调整几何 → 60ms 防抖后进入重算", "#FFF8E1", "#F9A825")
    y_roi_adj = y
    varrow(mid, y_roi + nd, y_roi_add, (249, 168, 37))
    varrow(mid, y_roi_add + nh, y_roi_adj, (249, 168, 37))
    d.text((L + nw + 8, y_roi + 18), "是", font=small, fill=(180, 120, 0))
    d.line([(L, y_roi + nd // 2), (bypass_x, y_roi + nd // 2)], fill=(180, 120, 0), width=2)
    d.text((bypass_x - 30, y_roi + 4), "否", font=small, fill=(180, 120, 0))

    y = y_roi_adj + nh + gap + 6
    phase(40, y, "【E 构建处理链】", "#EF6C00")
    y += 26
    d.line([(bypass_x, y_roi + nd // 2), (bypass_x, y - 4), (mid, y - 4)],
           fill=(180, 120, 0), width=2)
    draw_arrow(d, (mid - 1, y - 4), (mid, y - 4), (180, 120, 0), 2)

    e_steps = [
        "左侧拖入算法块（MIME = 插件稳定 id）",
        "PluginManager::createBlock(id)",
        "addBlockToPanel + ImageProcessor::addBlock",
        "调参 / 使能 / 标题栏拖拽换序 / 删除",
        "可选：导入链 JSON / 右键复制粘贴 / Ctrl+Z",
    ]
    e_ys = []
    for t in e_steps:
        node(L, y, nw, nh, t, "#FFF3E0", "#EF6C00")
        e_ys.append(y)
        y += nh + gap
    varrow(mid, y_roi_adj + nh, e_ys[0], (249, 168, 37))
    for i in range(len(e_ys) - 1):
        varrow(mid, e_ys[i] + nh, e_ys[i + 1], (239, 108, 0))

    y = e_ys[-1] + nh + gap + 6
    phase(40, y, "【F 统一重算 onApplyProcessing】", "#2E7D32")
    y += 26
    f_steps = [
        ("getAllRoiInfo → 收集场景中全部 ROI", "#E8F5E9", "#2E7D32"),
        ("setRois(list)；空列表语义 = 全图", "#E8F5E9", "#2E7D32"),
        ("current = original（不破坏原图缓存）", "#E8F5E9", "#2E7D32"),
        ("对每个启用块：current = process(...)", "#E8F5E9", "#2E7D32"),
        ("m_result = current；emit processingFinished(ms)", "#E8F5E9", "#2E7D32"),
        ("刷新画布显示结果 + 耗时/信息标签", "#E3F2FD", "#1565C0"),
    ]
    f_ys = []
    for t, fill, border in f_steps:
        node(L, y, nw, nh, t, fill, border)
        f_ys.append(y)
        y += nh + gap
    varrow(mid, e_ys[-1] + nh, f_ys[0], (239, 108, 0))
    for i in range(len(f_ys) - 1):
        varrow(mid, f_ys[i] + nh, f_ys[i + 1])

    y_end = f_ys[-1] + nh + gap
    node(L, y_end, nw, nh, "结束本轮 / 等待下一次用户操作", "#E8F5E9", "#2E7D32")
    varrow(mid, f_ys[-1] + nh, y_end)

    rx = 780
    rounded_rect(d, (rx, 80, rx + 500, 420), 10, "#Faf5ff", "#7B1FA2", 2)
    d.text((rx + 16, 92), "【G 后续操作】（任意时刻）", font=box_f, fill=(123, 31, 162))
    for i, t in enumerate([
        "按住对比：pressed=原图，released=结果",
        "保存当前结果图（路径/格式自选）",
        "导出 / 导入处理链 JSON",
        "清空整条处理链 → 结果复位原图",
        "换图前：saveCurrentSession（链+ROI）",
        "退出时：saveSessionsToDisk → app_sessions.json",
        "设置：中/英、浅/深主题、热添加插件 DLL",
        "热添加成功 → rebuildAlgoList 刷新列表",
    ]):
        d.text((rx + 18, 130 + i * 32), "· " + t, font=small, fill=(45, 45, 45))

    rounded_rect(d, (rx, 450, rx + 500, 720), 10, "#FFF8E1", "#EF6C00", 2)
    d.text((rx + 16, 462), "触发重算 → 回到【F】", font=box_f, fill=(230, 81, 0))
    for i, t in enumerate([
        "拖入 / 删除 / 换序处理块",
        "块参数变化 / 使能开关",
        "ROI 增删或几何变化（约 60ms 防抖）",
        "点击「应用」按钮",
        "导入处理链 / 恢复会话之后",
        "撤销（Ctrl+Z）恢复结构之后",
    ]):
        d.text((rx + 18, 500 + i * 32), "· " + t, font=small, fill=(45, 45, 45))

    rounded_rect(d, (rx, 750, rx + 500, 920), 10, "#E3F2FD", "#1565C0", 2)
    d.text((rx + 16, 762), "换图回环 → 回到【B】", font=box_f, fill=(21, 101, 192))
    for i, t in enumerate([
        "文件夹缩略图点击另一张图",
        "或菜单再次「打开图片/文件夹」",
        "先保存当前会话，再加载新图并走 C～F",
    ]):
        d.text((rx + 18, 800 + i * 32), "· " + t, font=small, fill=(45, 45, 45))

    d.line([(L + nw, f_ys[-1] + nh // 2), (rx, 200)], fill=(123, 31, 162), width=2)
    draw_arrow(d, (rx - 1, 200), (rx, 200), (123, 31, 162), 2)
    d.text((L + nw + 6, f_ys[-1]), "可进行", font=tiny, fill=(123, 31, 162))

    d.line([(rx, 600), (L + nw + 24, 600), (L + nw + 24, f_ys[0] + nh // 2),
            (L + nw, f_ys[0] + nh // 2)], fill=(230, 81, 0), width=2)
    draw_arrow(d, (L + nw + 1, f_ys[0] + nh // 2), (L + nw, f_ys[0] + nh // 2), (230, 81, 0), 2)
    d.text((rx - 68, 575), "再重算", font=tiny, fill=(230, 81, 0))

    d.line([(rx, 820), (bypass_x, 820), (bypass_x, y_open + nd // 2), (L, y_open + nd // 2)],
           fill=(21, 101, 192), width=2)
    draw_arrow(d, (L + 1, y_open + nd // 2), (L, y_open + nd // 2), (21, 101, 192), 2)
    d.text((bypass_x + 8, 790), "换图", font=tiny, fill=(21, 101, 192))

    d.text((40, 1925),
           "主路径自上而下（A→F）；右侧为后续操作、重算回环与换图回环。",
           font=tiny, fill=(90, 90, 90))

    path = FIG / "21_main_flow.png"
    img.save(path)
    return path

def make_class_diagram():
    w, h = 1500, 1040
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small = font(22, True), font(12, True), font(11)
    d.text((30, 12), "类图（插件接口 · 核心继承与协作）", font=title_f, fill=(33, 33, 33))

    def box(x, y, ww, hh, title, lines, fill, border):
        rounded_rect(d, (x, y, x + ww, y + hh), 8, fill, border, 2)
        d.text((x + 10, y + 8), title, font=box_f, fill=border)
        d.line([(x + 8, y + 30), (x + ww - 8, y + 30)], fill=border, width=1)
        for i, line in enumerate(lines):
            d.text((x + 10, y + 38 + i * 17), line, font=small, fill=(40, 40, 40))

    box(30, 45, 300, 200, "Widget : IBlockHost", [
        "主窗口 / 交互总控",
        "+ loadImageFromPath()",
        "+ onApplyProcessing()",
        "+ getAllRoiInfo()",
        "+ createBlockByName()",
        "+ onAddPlugin / 主题切换",
        "+ 会话落盘 / 撤销 / 对比",
    ], "#E3F2FD", "#1565C0")
    box(360, 45, 280, 200, "ImageProcessor", [
        "处理链调度引擎",
        "+ setOriginalImage / setRois",
        "+ add/remove/moveBlock",
        "+ resetResultToOriginal()",
        "+ reprocess()",
        "sig: finished / request",
    ], "#E8F5E9", "#2E7D32")
    box(670, 45, 280, 200, "PluginManager", [
        "<<singleton>>",
        "+ loadFromDirectory()",
        "+ loadPluginFile()",
        "+ createBlock(id)",
        "+ setEnglishUi()",
        "+ infos() / ids()",
    ], "#E8EAF6", "#283593")
    box(980, 45, 230, 130, "IBlockPlugin", [
        "<<interface>>",
        "+ id() / order()",
        "+ createBlock()",
    ], "#EDE7F6", "#5E35B1")
    box(1240, 45, 230, 130, "IBlockHost", [
        "<<interface>>",
        "+ hostHasImage()",
        "+ hostOriginalImage()",
        "+ hostCurrentRois()",
    ], "#EDE7F6", "#5E35B1")

    box(30, 280, 260, 150, "BaseBlock", [
        "<<abstract>> QWidget",
        "+ process(input, rois)*",
        "+ blockName()*",
        "+ saveParams/loadParams",
        "sig: params/remove/…",
    ], "#FFF3E0", "#EF6C00")

    blocks = [
        "Binarization", "Morphology", "Filter",
        "GrayTransform", "PseudoColor", "Glcm",
    ]
    for i, name in enumerate(blocks):
        x = 320 + (i % 3) * 230
        y = 270 + (i // 3) * 105
        extra = "直通图+特征面板" if name == "Glcm" else "UI + process()"
        box(x, y, 215, 88, name + "Block", ["extends BaseBlock", "in plugins/" + name.lower()],
            "#FFFDE7", "#F9A825")

    box(1020, 270, 220, 100, "ImageSession", [
        "<<struct>> chain+rois",
        "toJson / fromJson",
    ], "#E0F2F1", "#00695C")
    box(1260, 270, 200, 100, "AppLogger", [
        "<<static>>",
        "info/warn/error",
    ], "#F3E5F5", "#7B1FA2")

    box(30, 520, 200, 95, "ResizableRectItem", ["轴对齐矩形 ROI"], "#E8F5E9", "#43A047")
    box(250, 520, 220, 95, "ResizableEllipseItem", ["椭圆 ROI"], "#E8F5E9", "#43A047")
    box(490, 520, 250, 95, "ResizableRotatedRectItem", ["旋转矩形 ROI"], "#E8F5E9", "#43A047")
    box(760, 520, 200, 95, "RoiInfo", ["Rect/Ellipse/Rot", "toJson/fromJson"], "#FFF8E1", "#F9A825")
    box(980, 520, 200, 95, "RoiProcess", ["makeMask(list)", "apply 并集合成"], "#F3E5F5", "#7B1FA2")
    box(1200, 520, 260, 95, "ImageConverter / TimeMeasurer", ["QPixmap ↔ Mat", "reprocess 计时"], "#E0F7FA", "#00838F")

    box(30, 660, 520, 120, "plugins/*Plugin : IBlockPlugin", [
        "BinarizationPlugin / MorphologyPlugin / FilterPlugin …",
        "Q_PLUGIN_METADATA + createBlock → new *Block",
        "同目录 *Algorithm / Otsu / Glcm — 纯 OpenCV，无 Qt 主窗口依赖",
    ], "#ECEFF1", "#546E7A")
    box(570, 660, 420, 120, "AppConfig / StyleLoader", [
        "块名常量、MIME、主题色、大图上限",
        "loadTheme(light|dark) → QSS",
    ], "#F3E5F5", "#7B1FA2")
    box(1010, 660, 450, 120, "MyListWidget", [
        "继承 QListWidget",
        "mimeData → 块稳定 id（插件 infos）",
    ], "#E8EAF6", "#3949AB")

    draw_arrow(d, (330, 120), (360, 120), "#1565C0", 2)
    d.text((335, 95), "拥有", font=small, fill=(21, 101, 192))
    draw_arrow(d, (640, 145), (670, 145), "#283593", 2)
    d.text((645, 120), "工厂", font=small, fill=(40, 53, 147))
    draw_arrow(d, (810, 245), (1090, 175), "#5E35B1", 2)
    draw_arrow(d, (180, 245), (180, 280), "#EF6C00", 2)
    draw_arrow(d, (330, 175), (1240, 110), "#5E35B1", 1)
    for i in range(6):
        x = 320 + (i % 3) * 230 + 107
        y = 270 + (i // 3) * 105
        draw_arrow(d, (160, 355), (x, y), "#EF6C00", 1)

    d.text((30, 810), "关系要点：Widget 实现 IBlockHost 并 setHost；PluginManager 持有 IBlockPlugin 实例；",
          font=small, fill=(60, 60, 60))
    d.text((30, 835), "createBlockByName → PluginManager::createBlock；六种 Block 均由对应插件 DLL 创建。",
          font=small, fill=(60, 60, 60))

    path = FIG / "22_class.png"
    img.save(path)
    return path


def make_block_inherit_diagram():
    """算法链继承 + 插件接口"""
    w, h = 1400, 620
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small = font(22, True), font(13, True), font(11)
    d.text((28, 14), "算法链模块：IBlockPlugin → BaseBlock 继承关系（6 插件）", font=title_f, fill=(33, 33, 33))

    def box(x, y, ww, hh, title, lines, fill, border):
        rounded_rect(d, (x, y, x + ww, y + hh), 8, fill, border, 2)
        d.text((x + 12, y + 10), title, font=box_f, fill=border)
        for i, line in enumerate(lines):
            d.text((x + 12, y + 36 + i * 18), line, font=small, fill=(45, 45, 45))

    box(80, 50, 320, 90, "IBlockPlugin（接口）", [
        "id() · order() · createBlock()",
    ], "#EDE7F6", "#5E35B1")
    box(520, 50, 360, 90, "QWidget", ["Qt 控件基类"], "#ECEFF1", "#546E7A")
    box(460, 180, 480, 110, "BaseBlock（抽象，编入 blocksdk）", [
        "process(input, rois)*  ·  blockName()*",
        "saveParams / loadParams · 使能/删除 UI",
        "paramsChanged / remove / copy / paste …",
    ], "#FFF3E0", "#EF6C00")
    draw_arrow(d, (700, 140), (700, 180), "#546E7A", 2)
    draw_arrow(d, (240, 140), (560, 230), "#5E35B1", 2)
    d.text((300, 175), "插件 createBlock 返回", font=small, fill=(94, 53, 177))

    children = [
        ("BinarizationBlock", "阈值 + Otsu"),
        ("MorphologyBlock", "膨胀腐蚀开闭等"),
        ("FilterBlock", "均值/高斯/边缘"),
        ("GrayTransformBlock", "亮度对比度等"),
        ("PseudoColorBlock", "伪彩色映射"),
        ("GlcmBlock", "纹理特征直通"),
    ]
    for i, (name, desc) in enumerate(children):
        x = 20 + i * 225
        box(x, 360, 215, 100, name, [desc, "plugins/.../*.dll"], "#FFFDE7", "#F9A825")
        draw_arrow(d, (700, 290), (x + 107, 360), "#EF6C00", 2)

    d.text((28, 500), "每个插件目录同时包含：*Plugin（导出）+ *Block（UI）+ *Algorithm（OpenCV 纯函数）。",
          font=small, fill=(60, 60, 60))
    d.text((28, 530), "主程序不再内置算法实现；扩展方式：新 DLL 实现 IBlockPlugin，放入 exe/plugins/。",
          font=small, fill=(60, 60, 60))

    path = FIG / "23_block_inherit.png"
    img.save(path)
    return path


def make_plugin_dev_diagram():
    """开发文档用：新增插件步骤图"""
    w, h = 1200, 520
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small = font(22, True), font(13, True), font(12)
    d.text((28, 14), "新增算法插件开发流程", font=title_f, fill=(33, 33, 33))

    steps = [
        ("1. 建目录", "plugins/myalgo/\n.pro + 源码", "#E3F2FD", "#1565C0"),
        ("2. 实现算法", "*Algorithm\n纯 OpenCV", "#E0F7FA", "#00838F"),
        ("3. 实现块", "MyAlgoBlock\n: BaseBlock", "#FFF3E0", "#EF6C00"),
        ("4. 导出插件", "MyAlgoPlugin\n: IBlockPlugin", "#EDE7F6", "#5E35B1"),
        ("5. 登记工程", "plugins.pro\n加入 SUBDIRS", "#E8F5E9", "#2E7D32"),
        ("6. 编译运行", "生成 block_*.dll\n放入 plugins/", "#F3E5F5", "#7B1FA2"),
    ]
    for i, (title, body, fill, border) in enumerate(steps):
        x = 40 + i * 190
        rounded_rect(d, (x, 70, x + 170, 200), 10, fill, border, 2)
        d.text((x + 14, 82), title, font=box_f, fill=border)
        for j, line in enumerate(body.split("\n")):
            d.text((x + 14, 120 + j * 28), line, font=small, fill=(40, 40, 40))
        if i < len(steps) - 1:
            draw_arrow(d, (x + 170, 135), (x + 190, 135), (80, 80, 80), 2)

    notes = [
        "约定：TARGET = block_<name>；id() 与 blockName() 一致且稳定（勿随语言变）。",
        "链接：plugin_common.pri 已配置 blocksdk + OpenCV；可选拷贝 i18n/*_en.qm。",
        "验证：重启或「设置→添加插件」加载后，左侧列表出现新块并可拖入处理链。",
    ]
    for i, t in enumerate(notes):
        d.text((40, 300 + i * 36), t, font=small, fill=(50, 50, 50))

    path = FIG / "24_plugin_dev.png"
    img.save(path)
    return path


def make_dirs_diagram():
    """工程目录结构示意"""
    w, h = 1100, 720
    img = Image.new("RGB", (w, h), (255, 255, 255))
    d = ImageDraw.Draw(img)
    title_f, box_f, small = font(22, True), font(13, True), font(12)
    d.text((28, 14), "工程目录与产物布局", font=title_f, fill=(33, 33, 33))

    lines = [
        "opencv/",
        "├── opencv.pro                 # SUBDIRS: blocksdk → plugins → opencv_app",
        "├── main.cpp / widget.ui / resources.qrc",
        "├── core/                     # Widget · ImageProcessor · ImageSession",
        "├── blocks/                   # BaseBlock（编入 blocksdk）",
        "├── blocksdk/                 # IBlockPlugin · PluginManager → blocksdk.dll",
        "├── plugins/                  # 六算法源码 → bin/*/plugins/block_*.dll",
        "│   ├── binarization/ morphology/ filter/ …",
        "│   └── plugin_common.pri",
        "├── roi/  utils/  config/  styles/  i18n/",
        "├── docs/                     # 本文档与附图生成脚本",
        "└── bin/debug|release/",
        "    ├── opencv.exe",
        "    ├── blocksdk.dll",
        "    ├── plugins/block_*.dll (+ *_en.qm)",
        "    ├── sessions/app_sessions.json",
        "    └── logs/app_yyyyMMdd.log",
    ]
    rounded_rect(d, (40, 55, 1060, 680), 12, "#F8FAFC", "#64748B", 2)
    for i, line in enumerate(lines):
        d.text((60, 70 + i * 34), line, font=small, fill=(30, 30, 30))

    path = FIG / "25_dirs_layout.png"
    img.save(path)
    return path


# ========================= Word 工具 =========================

def set_run_font(run, size=11, bold=False, name="微软雅黑", color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_bookmark(paragraph, bookmark_name, bookmark_id):
    tag = paragraph._p
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    tag.insert(0, start)
    tag.append(end)


def add_heading(doc, text, level=1, bookmark_name=None, bookmark_id=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else (13 if level == 2 else 12),
                     bold=True, name="微软雅黑")
    if bookmark_name is not None and bookmark_id is not None:
        add_bookmark(p, bookmark_name, bookmark_id)
    return p


def add_para(doc, text, size=11, bold=False, center=False, space_after=6):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_code(doc, text, size=9):
    """等宽代码/伪代码块（上报文档中的关键片段）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, name="Consolas")
    # 浅灰底提示为代码区
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F5F5")
    shd.set(qn("w:val"), "clear")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        set_cell_shading(cell, "0F766E")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                set_cell_shading(cell, "F5F5F5")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=10, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(10)


def add_picture(doc, path, width_in=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))


def add_hyperlink_to_bookmark(paragraph, text, bookmark_name):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), bookmark_name)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "24")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "微软雅黑")
    rFonts.set(qn("w:hAnsi"), "微软雅黑")
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rPr.append(rFonts)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_toc_entry(doc, title, bookmark, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent = Cm((level - 1) * 0.75)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(14.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    add_hyperlink_to_bookmark(p, title, bookmark)
    return p


def enable_update_fields_on_open(doc):
    settings = doc.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


# ========================= 正文 =========================

def build_doc(figures):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)
    enable_update_fields_on_open(doc)

    bm = {
        "ch1": ("bm_ch1", 1),
        "s11": ("bm_s11", 2),
        "s12": ("bm_s12", 3),
        "s13": ("bm_s13", 4),
        "ch2": ("bm_ch2", 5),
        "s21": ("bm_s21", 6),
        "s22": ("bm_s22", 7),
        "s221": ("bm_s221", 8),
        "s222": ("bm_s222", 9),
        "s223": ("bm_s223", 10),
        "s224": ("bm_s224", 11),
        "s225": ("bm_s225", 12),
        "s226": ("bm_s226", 13),
        "s227": ("bm_s227", 14),
        "s23": ("bm_s23", 15),
        "s24": ("bm_s24", 16),
        "ch3": ("bm_ch3", 17),
        "s31": ("bm_s31", 18),
        "s32": ("bm_s32", 19),
        "s33": ("bm_s33", 20),
        "s34": ("bm_s34", 21),
        "s35": ("bm_s35", 22),
        "s36": ("bm_s36", 23),
        "s37": ("bm_s37", 24),
        "app": ("bm_app", 25),
    }

    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, "图像处理工具", size=28, bold=True, center=True, space_after=12)
    add_para(doc, "需求与设计说明", size=20, bold=True, center=True, space_after=10)
    add_para(doc, "技术栈：C++17 / Qt 6 Widgets / OpenCV 4.x / qmake（插件化）", size=11, center=True)
    add_para(doc, "工程名：opencv　　版本：1.0.0　　文档日期：2026-07-28", size=11, center=True, space_after=16)
    add_para(doc,
             "本文结构：第一章给出功能/性能/DFX 需求；"
             "第二章给出概要与详细设计（模块框图、模块说明、流程图、类图）；"
             "第三章给出开发文档（分层设计动机、接口职责、插件扩展完整示例）。"
             "内容与当前完整实现同步（插件化、多 ROI、GLCM、撤销、会话落盘、浅/深主题、运行日志）。",
             size=10, center=True)
    doc.add_page_break()

    add_para(doc, "目录", size=18, bold=True, center=True, space_after=12)
    add_para(doc, "（点击条目可跳转到对应章节）", size=10, center=True, space_after=10)

    toc_items = [
        ("第一章 需求", "bm_ch1", 1),
        ("1.1 功能需求", "bm_s11", 2),
        ("1.2 性能需求", "bm_s12", 2),
        ("1.3 DFX 需求", "bm_s13", 2),
        ("第二章 概要设计与详细设计", "bm_ch2", 1),
        ("2.1 总体结构与模块框图", "bm_s21", 2),
        ("2.2 模块说明", "bm_s22", 2),
        ("2.2.1 主界面模块 Widget", "bm_s221", 3),
        ("2.2.2 处理调度模块 ImageProcessor", "bm_s222", 3),
        ("2.2.3 算法链模块 BaseBlock 族", "bm_s223", 3),
        ("2.2.4 ROI 模块", "bm_s224", 3),
        ("2.2.5 会话与撤销", "bm_s225", 3),
        ("2.2.6 插件框架与算法库", "bm_s226", 3),
        ("2.2.7 配置、样式、国际化与日志", "bm_s227", 3),
        ("2.3 流程图", "bm_s23", 2),
        ("2.4 类图", "bm_s24", 2),
        ("第三章 开发文档", "bm_ch3", 1),
        ("3.1 本章说明与技术基线", "bm_s31", 2),
        ("3.2 分层架构与设计动机", "bm_s32", 2),
        ("3.3 核心接口职责", "bm_s33", 2),
        ("3.4 运行时协作关系", "bm_s34", 2),
        ("3.5 新增算法插件完整示例", "bm_s35", 2),
        ("3.6 处理与数据约定", "bm_s36", 2),
        ("3.7 工程交付与维护要点", "bm_s37", 2),
        ("附录 术语", "bm_app", 1),
    ]
    for title, bookmark, level in toc_items:
        add_toc_entry(doc, title, bookmark, level)

    doc.add_page_break()

    # =====================================================================
    # 第一章
    # =====================================================================
    add_heading(doc, "第一章 需求", 1, *bm["ch1"])
    add_para(doc,
             "本章描述本工具应满足的功能需求、性能需求与常规 DFX 需求。"
             "需求表统一三列：编号、需求描述、说明与规格。")

    add_heading(doc, "1.1 功能需求", 2, *bm["s11"])
    add_table(doc, ["编号", "需求描述", "说明与规格"], [
        ["FR-01", "打开单张图片并在画布显示", "支持 png/jpg/jpeg/bmp/gif/tiff"],
        ["FR-02", "打开文件夹并以缩略图浏览切换", "点击缩略图切换当前图；每张图独立会话；切换图片时自动保存会话"],
        ["FR-03", "画布缩放与平移",
         "【缩放】范围 0.1～5.0，滚轮按倍数步长缩放（锚点为鼠标下场景点）；"
         "【平移】中键或空白处拖拽平滑平移；【复位】Ctrl+0 适应窗口完整显示"],
        ["FR-04", "显示图像信息与处理耗时", "宽高、缩放比例、文件名；重算后显示毫秒耗时"],
        ["FR-05", "从算法列表拖拽创建处理块", "左侧工具箱由已加载插件动态填充；拖到右侧处理链面板"],
        ["FR-06", "处理链按顺序流水线执行", "仅执行「启用」块；顺序即执行顺序"],
        ["FR-07", "二值化处理（含 Otsu）",
         "灰度∈[下限,上限]为白；灰度下限/上限取值范围均为 0～255；一键 Otsu 写回阈值"],
        ["FR-08", "形态学处理",
         "膨胀/腐蚀/开/闭/顶帽/底帽/梯度；核大小须为奇数，默认 3×3，x/y 最大不超过 31；"
         "迭代次数为整数，范围 1～20"],
        ["FR-09", "滤波处理",
         "均值/高斯/中值及 Sobel/Laplacian/Prewitt/Roberts；核大小须为奇数，默认 3×3，"
         "x/y 最大不超过 31；迭代次数为整数，范围 1～20"],
        ["FR-10", "灰度变换", "转灰、亮度对比度、反色、对数、伽马、均衡化、归一化；默认转灰度图"],
        ["FR-11", "伪彩色映射", "Jet/Hot/Cool/Rainbow 等 10 种色带"],
        ["FR-12", "灰度共生矩阵（GLCM）分析", "图像直通；面板显示对比度/相关/能量/均匀性/熵/相异性"],
        ["FR-13", "处理块可调参、使能、删除", "参数变化触发重算；可临时关闭某块"],
        ["FR-14", "处理块可拖拽换序", "右侧面板内拖动改变执行顺序"],
        ["FR-15", "处理块可复制/粘贴参数", "右键菜单；剪贴板存单块 JSON；面板空白处可粘贴追加"],
        ["FR-16", "清空整条处理链", "确认后移除所有块，结果恢复为原图"],
        ["FR-17", "多 ROI：矩形/椭圆/旋转矩形", "可同时存在多个；Delete 删选中或全部"],
        ["FR-18", "多 ROI 并集局部处理", "区域外保持原像素；空列表=全图；多个 ROI 重叠时按并集处理"],
        ["FR-19", "按住对比原图与结果图", "按下看原图，松开恢复结果"],
        ["FR-20", "保存当前结果图", "用户选择路径与格式（png/jpeg/bmp 等）"],
        ["FR-21", "处理链导入/导出为 JSON", "导出含块类型与参数；导入替换现有链；导入失败时保留旧链"],
        ["FR-22", "按图片路径记忆处理链与 ROI", "换图自动保存/恢复；退出写入 sessions/app_sessions.json"],
        ["FR-23", "撤销结构变更（Ctrl+Z）",
         "【撤销范围】仅针对结构变更（增删块、换序、增删 ROI、清空、导入 JSON）；"
         "【深度限制】最多记忆 40 步"],
        ["FR-24", "界面中/英切换", "设置菜单切换；语言偏好可持久化；插件自带 *_en.qm 同步切换"],
        ["FR-25", "运行日志与关于/帮助", "日志写入 exe/logs/；关于可查看日志与会话路径；快捷键帮助"],
        ["FR-26", "浅色 / 深色主题切换", "设置菜单切换；StyleLoader 加载 theme_light/dark.qss；偏好持久化"],
        ["FR-27", "算法以插件 DLL 方式加载",
         "启动扫描 exe/plugins/block_*.dll；主程序空壳不内置算法实现；"
         "未加载到插件时提示正确运行路径"],
        ["FR-28", "运行时热添加算法插件", "设置→添加插件，选择 DLL 后注册并刷新左侧算法列表"],
        ["FR-29", "支持大图加载", "取消 Qt6 QImageReader 默认 256MB 分配上限，仍受系统内存约束"],
    ], [2.2, 5.8, 7.2])

    add_heading(doc, "1.2 性能需求", 2, *bm["s12"])
    add_table(doc, ["编号", "需求描述", "说明与规格"], [
        ["PR-01", "调参后自动重算并尽快反馈", "常规分辨率（约 1920×1080）交互可接受"],
        ["PR-02", "ROI 拖动时防抖重算", "几何变化经约 60ms 防抖后再重算，避免拖动卡顿"],
        ["PR-03", "重算耗时可观测", "界面显示最近一次 reprocess 毫秒"],
        ["PR-04", "缩放浏览不触发算法链", "仅改变视图变换矩阵"],
        ["PR-05", "撤销栈有上限", "单图最多保留 40 层结构快照，避免内存无限增长"],
        ["PR-06", "插件启动加载可感知", "日志记录尝试目录与加载个数；失败弹窗提示运行路径"],
    ], [2.2, 5.8, 7.2])

    add_heading(doc, "1.3 DFX 需求", 2, *bm["s13"])
    add_para(doc, "本节采用常规、简化的 DFX（可靠性、可维护性、易用性、可扩展性、兼容性）。")
    add_table(doc, ["编号", "需求描述", "说明与规格"], [
        ["DX-01", "可靠性：空图/非法文件不崩溃", "打开失败提示；无图时重算安全返回"],
        ["DX-02", "可靠性：参数越界由控件约束", "SpinBox/组合框限制合法范围"],
        ["DX-03", "可靠性：关键操作可追踪", "AppLogger 记录打开、保存、导入导出、插件加载、异常等"],
        ["DX-04", "可维护性：主程序与算法插件分层",
         "算法与块 UI 位于 plugins/；BaseBlock/PluginManager 位于 blocksdk；主程序负责交互与调度"],
        ["DX-05", "可扩展性：新增算法块有固定步骤",
         "实现 IBlockPlugin + BaseBlock 子类，编译为 DLL 放入 plugins/（或热添加）即可扩展"],
        ["DX-06", "易用性：拖拽建链、按住对比、撤销、主题", "降低调参与试错成本；浅/深主题适配不同环境"],
        ["DX-07", "易用性：中英界面与快捷键说明", "降低上手成本"],
        ["DX-08", "兼容性：Windows + Qt6 + OpenCV4", "以当前 qmake 工程为交付基线"],
    ], [2.2, 5.8, 7.2])

    doc.add_page_break()

    # =====================================================================
    # 第二章
    # =====================================================================
    add_heading(doc, "第二章 概要设计与详细设计", 1, *bm["ch2"])
    add_para(doc,
             "本章在需求基础上给出模块划分、数据流与接口，说明各模块职责与实现要点，"
             "并给出主流程图与类图。架构已切换为「主程序空壳 + blocksdk + 算法插件」。")

    add_heading(doc, "2.1 总体结构与模块框图", 2, *bm["s21"])
    add_para(doc,
             "系统按「表现层 → 调度层 → 插件块层 → 算法/工具/会话层」组织。"
             "Widget 负责交互、多 ROI、撤销、会话、主题与热添加插件；"
             "PluginManager（blocksdk）负责扫描/注册/创建块；"
             "ImageProcessor 按链调度；各插件内 BaseBlock 子类负责参数 UI 与单步处理；"
             "utils / RoiProcess / AppLogger 提供计算与基础设施。")
    add_picture(doc, figures["block"], 6.4)
    add_caption(doc, "图 2-1 总体模块框图（插件化 · 含主要数据流与接口）")

    add_para(doc, "主要接口约定：", bold=True)
    add_table(doc, ["接口/数据", "提供方", "消费方 / 说明"], [
        ["QPixmap 原图/结果图", "ImageProcessor", "Widget 显示；链输入输出"],
        ["QList<RoiInfo>", "Widget::getAllRoiInfo", "setRois → Block::process；空=全图"],
        ["BaseBlock*", "PluginManager::createBlock", "面板布局 + ImageProcessor::addBlock"],
        ["process(input, rois)", "各 *Block（插件）", "reprocess 顺序调用"],
        ["IBlockPlugin", "plugins/*.dll", "PluginManager 加载与注册"],
        ["IBlockHost", "Widget", "块取原图/ROI（如 Otsu）"],
        ["paramsChanged", "BaseBlock", "→ requestReprocess"],
        ["requestReprocess", "ImageProcessor", "Widget 先同步 ROI 再 reprocess"],
        ["processingFinished(ms)", "ImageProcessor", "刷新画布与耗时"],
        ["ImageSession", "Widget", "换图记忆；撤销快照；落盘 JSON"],
        ["主题 QSS", "StyleLoader", "qApp->setStyleSheet"],
        ["日志事件", "各模块", "AppLogger → exe/logs/"],
    ], [4.5, 4.5, 6.5])

    add_heading(doc, "2.2 模块说明", 2, *bm["s22"])

    add_heading(doc, "2.2.1 主界面模块 Widget", 3, *bm["s221"])
    add_para(doc,
             "位置：core/widget.h/.cpp + widget.ui。"
             "职责：文件打开与文件夹缩略图、画布、多 ROI 图元、处理链面板、"
             "菜单（文件含撤销；设置含语言、主题、添加插件）、对比/保存/链 JSON、"
             "会话与撤销、中英切换。"
             "实现 IBlockHost，供插件块查询原图与当前 ROI。"
             "统一重算入口 onApplyProcessing：getAllRoiInfo → setRois → reprocess。"
             "createBlockByName 委托 PluginManager::createBlock；"
             "rebuildAlgoList 按已注册插件刷新左侧列表。"
             "相关类：MyListWidget、三种 Resizable*Item、ImageSession、PluginManager。")

    add_heading(doc, "2.2.2 处理调度模块 ImageProcessor", 3, *bm["s222"])
    add_para(doc,
             "位置：core/imageprocessor.h/.cpp，无界面。"
             "持有 m_original、m_result、m_blocks、m_rois（列表）。"
             "结构变化只发 requestReprocess；reprocess 内 TimeMeasurer 计时。"
             "resetResultToOriginal 用于清空链后把结果复位为原图。"
             "与插件解耦：只依赖 BaseBlock 指针，不关心块来自哪个 DLL。")

    add_heading(doc, "2.2.3 算法链模块 BaseBlock 族", 3, *bm["s223"])
    add_para(doc,
             "位置：基类 blocks/baseblock.*（编入 blocksdk）；"
             "六子类位于 plugins/<name>/。"
             "BaseBlock 提供标题栏、使能/删除、拖拽换序、右键复制粘贴，"
             "纯虚接口 process(input, rois) / blockName。"
             "六插件块：Binarization / Morphology / Filter / GrayTransform / PseudoColor / Glcm。"
             "GLCM 为分析块：图像原样通过，面板展示纹理统计（量化级、距离可调）。")
    add_picture(doc, figures["inherit"], 6.3)
    add_caption(doc, "图 2-2 算法链 BaseBlock 继承关系（含插件接口）")
    add_table(doc, ["类", "父类/关系", "要点"], [
        ["BaseBlock", "QWidget", "抽象块；UI 骨架与信号；编入 blocksdk"],
        ["BinarizationBlock", "BaseBlock", "上下限 + Otsu；plugins/binarization"],
        ["MorphologyBlock", "BaseBlock", "Op / 核 / 迭代"],
        ["FilterBlock", "BaseBlock", "滤波类型与核"],
        ["GrayTransformBlock", "BaseBlock", "变换类型与亮度等"],
        ["PseudoColorBlock", "BaseBlock", "色带 Map"],
        ["GlcmBlock", "BaseBlock", "直通图 + 特征显示"],
        ["*Plugin", "IBlockPlugin", "Q_PLUGIN_METADATA 导出；createBlock"],
        ["ImageProcessor", "聚合块指针", "按序 process"],
        ["Widget", "创建并拥有块", "经 PluginManager 工厂创建"],
    ], [4, 4, 7.5])

    add_heading(doc, "2.2.4 ROI 模块", 3, *bm["s224"])
    add_para(doc,
             "位置：roi/ 与 utils/roiprocess.*"
             "RoiInfo 描述 Rect / Ellipse / RotatedRect，支持 toJson/fromJson。"
             "场景中可同时存在多个 Resizable*Item；getAllRoiInfo 收集为列表。"
             "RoiProcess::makeMask/apply 对列表做蒙版并集（OR）；"
             "先做全图处理再经蒙版只合成 ROI 区域内结果，区域外保持原像素；空列表=全图处理。"
             "几何变更经约 60ms 防抖后触发重算。")

    add_heading(doc, "2.2.5 会话与撤销", 3, *bm["s225"])
    add_para(doc,
             "位置：core/imagesession.h。"
             "ImageSession = 处理链 JSON 数组 + ROI 列表。"
             "Widget::m_sessions 以图片绝对路径为键；换图时保存当前、恢复目标；"
             "退出写入 {exe}/sessions/app_sessions.json（version=1）。"
             "撤销：结构变更前 pushUndoSnapshot，Ctrl+Z 弹出恢复；"
             "栈深度 MAX_UNDO=40；换图清空该图撤销栈。"
             "ROI 几何微调走防抖重算，不把每一帧拖动都压入撤销栈"
             "（压栈点为添加/删除 ROI 与链结构变更）。")

    add_heading(doc, "2.2.6 插件框架与算法库", 3, *bm["s226"])
    add_para(doc,
             "位置：blocksdk/（iblockplugin.h、iblockhost.h、pluginmanager.*）与 plugins/。"
             "PluginManager 单例：loadFromDirectory 扫描 DLL；loadPluginFile 热加载；"
             "createBlock(id) 调用已注册工厂；setEnglishUi 安装/卸载插件自带翻译。"
             "每个算法插件目录包含：*Plugin（导出）+ *Block（UI）+ *Algorithm（纯 OpenCV）。"
             "现有六插件与 order：二值化(10)、形态学(20)、滤波(30)、灰度(40)、伪彩(50)、GLCM(60)。"
             "utils/：ImageConverter、TimeMeasurer、RoiProcess、AppLogger（供主程序与插件共用）。"
             "Block::process 典型骨架：Pixmap→Mat → RoiProcess::apply(算法) → Mat→Pixmap；"
             "GlcmBlock 额外 compute 特征后仍返回原图。")

    add_heading(doc, "2.2.7 配置、样式、国际化与日志", 3, *bm["s227"])
    add_para(doc,
             "AppConfig：块名常量、默认阈值、缩放极限、MIME、主题色、大图分配上限。"
             "StyleLoader + styles/theme_light.qss / theme_dark.qss；设置菜单切换并写入 QSettings。"
             "i18n：主程序 opencv_en.qm + 各插件 *_en.qm；块列表用稳定 id 抗语言切换。"
             "AppLogger：exe/logs/app_yyyyMMdd.log；关于对话框可提示路径。"
             "扩展新块：见第三章「新增算法插件步骤」。")

    add_heading(doc, "2.3 流程图", 2, *bm["s23"])
    add_para(doc, "下图描述启动加载插件、打开图像、可选多 ROI、建链、统一重算到显示/落盘的主路径。")
    add_picture(doc, figures["flow"], 5.8)
    add_caption(doc, "图 2-3 主流程图")

    add_para(doc, "重算核心步骤：", bold=True)
    add_table(doc, ["步骤", "动作", "说明"], [
        ["1", "getAllRoiInfo", "空列表 = 全图"],
        ["2", "setRois(list)", "写入本次重算 ROI"],
        ["3", "current = original", "不破坏 m_original"],
        ["4", "启用块依次 process", "current = block->process(current, rois)"],
        ["5", "m_result = current", "emit processingFinished"],
        ["6", "Widget 刷新画布", "更新耗时与信息标签"],
    ], [2, 5, 8.5])

    add_heading(doc, "2.4 类图", 2, *bm["s24"])
    add_para(doc,
             "下图汇总核心类协作：Widget 实现 IBlockHost 并拥有 ImageProcessor、会话与撤销栈；"
             "PluginManager 管理 IBlockPlugin；六种 Block 继承 BaseBlock 并由插件创建。")
    add_picture(doc, figures["class"], 6.4)
    add_caption(doc, "图 2-4 核心类图")

    add_table(doc, ["类/单元", "层", "职责摘要"], [
        ["Widget", "表现", "交互总控、多 ROI、链 UI、会话、撤销、语言/主题/热插件"],
        ["MyListWidget", "表现", "算法列表拖出稳定 id"],
        ["ImageProcessor", "应用", "处理链调度与结果缓存"],
        ["PluginManager", "应用/基础设施", "插件扫描、注册、创建、翻译"],
        ["IBlockPlugin / IBlockHost", "接口", "插件契约 / 宿主能力"],
        ["BaseBlock 及 6 子类", "应用+表现（插件）", "参数 UI + process"],
        ["ImageSession", "应用", "单图链+ROI 快照与落盘格式"],
        ["RoiInfo / Resizable*Item", "领域/表现", "ROI 数据与可视化编辑"],
        ["RoiProcess", "领域", "多 ROI 并集蒙版与局部合成"],
        ["*Algorithm / Glcm", "领域（插件内）", "OpenCV 算法与纹理特征"],
        ["ImageConverter / TimeMeasurer", "基础", "格式转换与计时"],
        ["AppConfig / StyleLoader / AppLogger", "基础", "配置、浅深主题、文件日志"],
    ], [4.8, 2.8, 7.9])

    doc.add_page_break()

    # =====================================================================
    # 第三章 开发文档（上报口径：分层动机 + 接口职责 + 插件扩展示例）
    # =====================================================================
    add_heading(doc, "第三章 开发文档", 1, *bm["ch3"])
    add_para(doc,
             "本章从工程实现角度说明：为何采用「主程序空壳 + SDK + 算法插件」分层、"
             "各核心接口的职责边界、运行时如何协作，以及如何在不改动主程序业务逻辑的前提下"
             "完整扩展一个新算法插件。"
             "表述侧重设计意图与可复核的扩展路径，便于技术评审与后续维护交接。")

    # ----- 3.1 -----
    add_heading(doc, "3.1 本章说明与技术基线", 2, *bm["s31"])
    add_para(doc,
             "本工具定位为桌面端图像处理工作台：用户通过拖拽组装处理链，对单图或多图会话"
             "进行交互式调参，并支持 ROI 局部处理、结果对比与会话记忆。"
             "算法能力以动态库插件形式交付，主程序负责交互与调度，二者通过稳定接口解耦。")
    add_table(doc, ["基线项", "说明"], [
        ["平台", "Windows 10/11 x64"],
        ["语言 / 框架", "C++17；Qt 6 Widgets；OpenCV 4.x（world 库）"],
        ["构建", "qmake；根工程 opencv.pro 为 SUBDIRS（ordered）"],
        ["版本", "应用版本 1.0.0（见 AppConfig::APP_VERSION）"],
        ["交付形态", "opencv.exe + blocksdk.dll + plugins/block_*.dll"],
    ], [3.5, 11.5])
    add_para(doc,
             "OpenCV 头文件与库路径集中配置于 opencv_app.pro 与 plugins/plugin_common.pri；"
             "换机部署时两处须同步修改并全量重建，以保证主程序与插件 ABI、链接库一致。")

    # ----- 3.2 -----
    add_heading(doc, "3.2 分层架构与设计动机", 2, *bm["s32"])
    add_para(doc,
             "若将全部算法编译进主程序，每次增减算法都要改主工程、拉长编译与发布周期，"
             "且算法实现与界面逻辑容易缠绕，不利于责任划分与回归。"
             "因此本工程采用三层产物分离：")
    add_table(doc, ["层次", "工程 / 产物", "职责", "设计动机"], [
        ["表现与编排层", "opencv_app → opencv.exe",
         "主窗口、画布、ROI、处理链面板、会话/撤销、主题与语言",
         "把「怎么用」集中在一处；不内嵌具体算法实现"],
        ["插件 SDK 层", "blocksdk → blocksdk.dll",
         "BaseBlock、PluginManager、IBlockPlugin / IBlockHost、共用工具",
         "主程序与插件共享同一套块模型与加载协议，避免两套实现"],
        ["算法插件层", "plugins → block_*.dll",
         "各算法的 UI 块 + 纯 OpenCV 计算 + Qt 插件导出类",
         "算法可独立编译、替换、热添加；主程序发版可不因算法增删而改业务代码"],
    ], [2.8, 4.2, 4.5, 4])
    add_picture(doc, figures["dirs"], 5.8)
    add_caption(doc, "图 3-1 工程目录与产物布局（与分层对应）")

    add_para(doc, "分层带来的直接收益：", bold=True)
    add_para(doc,
             "① 可扩展：新算法按插件契约交付即可接入工具箱与处理链；"
             "② 可维护：算法缺陷修复通常只需替换对应 DLL；"
             "③ 可测试：纯算法（*Algorithm）可脱离 UI 做单元级验证；"
             "④ 边界清晰：调度器 ImageProcessor 只认 BaseBlock*，不感知具体算法类型。")

    # ----- 3.3 -----
    add_heading(doc, "3.3 核心接口职责", 2, *bm["s33"])
    add_para(doc,
             "插件化能否成立，取决于接口是否稳定、职责是否单一。"
             "下列接口构成本工程扩展面的「合同」。")

    add_para(doc, "（1）IBlockPlugin — 插件导出契约", bold=True)
    add_para(doc,
             "每个算法 DLL 通过 Qt 插件机制导出一个实现类。"
             "主程序只依赖 IID（com.opencvlab.IBlockPlugin/1.0），不依赖具体插件头文件。")
    add_table(doc, ["方法", "职责"], [
        ["id()", "稳定标识；须与块的 blockName() 一致，用于列表、导入导出、会话还原"],
        ["order()", "左侧工具箱排序权重；数值越小越靠前"],
        ["createBlock(parent)", "工厂方法：创建该算法对应的 BaseBlock 子类实例"],
    ], [4.5, 10.5])
    add_code(doc,
             "// 伪代码：插件必须满足的最小合同\n"
             "class XxxPlugin : QObject, IBlockPlugin {\n"
             "    id()         → 稳定字符串（勿随语言翻译而改变）\n"
             "    order()      → 整数排序键\n"
             "    createBlock()→ new XxxBlock(parent)\n"
             "};")

    add_para(doc, "（2）IBlockHost — 宿主能力（反向依赖倒置）", bold=True)
    add_para(doc,
             "处理链中 process(input, rois) 的 input 是上一块输出（中间结果）。"
             "当块需要「原图」或「当前场景 ROI」时（例如二值化一键 Otsu），"
             "不能也不应直接依赖 Widget 类型。"
             "IBlockHost 由主窗口实现，块通过 BaseBlock::host() 访问，从而保持插件与主程序解耦。")
    add_table(doc, ["方法", "职责", "典型用途"], [
        ["hostHasImage()", "是否已加载图像", "操作前校验，避免空图计算"],
        ["hostOriginalImage()", "返回原图像素", "Otsu 等必须基于原图统计的场景"],
        ["hostCurrentRois()", "返回当前 ROI 列表", "按 ROI 外接区域收窄统计范围"],
    ], [4, 5, 6.5])

    add_para(doc, "（3）BaseBlock — 处理块统一抽象", bold=True)
    add_para(doc,
             "所有可视算法块继承 BaseBlock。基类负责标题栏、使能、删除、拖拽换序、"
             "复制粘贴与参数变更信号；子类只实现差异部分。")
    add_table(doc, ["类别", "API / 信号", "职责"], [
        ["必须实现", "process(input, rois) / blockName()", "单步处理；稳定 id"],
        ["建议重写", "saveParams() / loadParams() / retranslateUi()",
         "参数持久化与界面语言"],
        ["基类提供", "setupTitle / contentLayout / trackParamWidget / setHost",
         "搭 UI、登记改参控件、注入宿主"],
        ["对外信号", "paramsChanged / paramsAboutToChange / removeRequested / …",
         "驱动重算、撤销压栈、面板增删"],
    ], [2.8, 5.5, 7.2])

    add_para(doc, "（4）PluginManager — 注册表与工厂", bold=True)
    add_para(doc,
             "单例。启动时 loadFromDirectory 扫描 plugins/*.dll；"
             "运行期可用 loadPluginFile 热添加。"
             "createBlock(id) 查找工厂、构造块，并注入 IBlockHost。"
             "setEnglishUi 统一安装/卸载各插件自带的 *_en.qm。")

    add_para(doc, "（5）ImageProcessor — 流水线调度", bold=True)
    add_para(doc,
             "持有原图、结果图、块指针列表与 ROI 列表。"
             "结构变化只发 requestReprocess；真正计算在 reprocess 中按序调用启用块的 process。"
             "调度层不包含任何具体算法分支，保证「加插件不必改调度器」。")

    # ----- 3.4 -----
    add_heading(doc, "3.4 运行时协作关系", 2, *bm["s34"])
    add_para(doc, "从启动到一次重算，关键协作可概括为：")
    add_table(doc, ["阶段", "参与方", "协作要点"], [
        ["启动", "main → PluginManager → Widget",
         "扫描插件；setHost(this)；rebuildAlgoList 填左侧列表"],
        ["建块", "拖放 → createBlockByName → PluginManager::createBlock",
         "按稳定 id 工厂创建；块已带 host"],
        ["入链", "Widget 面板 + ImageProcessor::addBlock",
         "UI 顺序与调度链顺序保持一致"],
        ["重算", "onApplyProcessing",
         "getAllRoiInfo → setRois → reprocess → 刷新画布与耗时"],
        ["持久化", "ImageSession / 链 JSON",
         "块 saveParams 的 name 字段必须能再次 createBlock"],
    ], [2.5, 5, 8])
    add_code(doc,
             "// 伪代码：一次统一重算\n"
             "rois = Widget.getAllRoiInfo()          // 空 = 全图\n"
             "Processor.setRois(rois)\n"
             "current = Processor.originalImage()\n"
             "for block in Processor.blocks():\n"
             "    if block.enabled:\n"
             "        current = block.process(current, rois)\n"
             "Processor.result = current\n"
             "emit processingFinished(elapsedMs)\n"
             "Widget.refreshDisplay()")
    add_para(doc,
             "创建块时注入宿主（实现摘录）：")
    add_code(doc,
             "BaseBlock *PluginManager::createBlock(const QString &id, QWidget *parent)\n"
             "{\n"
             "    Entry *e = findEntry(id);\n"
             "    if (!e || !e->factory) return nullptr;\n"
             "    BaseBlock *block = e->factory(parent);\n"
             "    if (block) block->setHost(m_host);   // 注入 IBlockHost\n"
             "    return block;\n"
             "}")

    # ----- 3.5 -----
    add_heading(doc, "3.5 新增算法插件完整示例", 2, *bm["s35"])
    add_para(doc,
             "以下以扩展一个名为「示例增强」的算法为例，说明完整落地路径。"
             "目标：主程序零业务改动（至多增加块名常量），新 DLL 放入 plugins/ 即可出现在工具箱并参与处理链。")
    add_picture(doc, figures["plugin_dev"], 6.2)
    add_caption(doc, "图 3-2 新增算法插件开发流程")

    add_para(doc, "3.5.1 目录与工程登记", bold=True)
    add_table(doc, ["项", "约定"], [
        ["目录", "plugins/demoenhance/"],
        ["产物名", "block_demoenhance.dll（PLUGIN_TARGET = block_demoenhance）"],
        ["文件", "demoenhance.pro；demoenhanceplugin.*；demoenhanceblock.*；demoenhance.*（纯算法）"],
        ["登记", "plugins/plugins.pro 的 SUBDIRS 增加 demoenhance"],
        ["公共构建", "include(../plugin_common.pri)（已配置 blocksdk、OpenCV、输出目录）"],
        ["排序 order", "建议 ≥ 70，避免打乱现有 10～60 内置顺序"],
    ], [3, 12])

    add_para(doc, "3.5.2 插件导出类（IBlockPlugin）", bold=True)
    add_para(doc, "职责：向宿主声明「我是谁、怎么造块」。实现形态与现有二值化插件一致：")
    add_code(doc,
             "// demoenhanceplugin.h（结构示意）\n"
             "class DemoEnhancePlugin : public QObject, public IBlockPlugin {\n"
             "    Q_OBJECT\n"
             "    Q_PLUGIN_METADATA(IID IBlockPlugin_iid)\n"
             "    Q_INTERFACES(IBlockPlugin)\n"
             "public:\n"
             "    QString id() const override;                 // 稳定 id\n"
             "    int order() const override { return 70; }\n"
             "    BaseBlock *createBlock(QWidget *parent) override;\n"
             "};\n"
             "\n"
             "// demoenhanceplugin.cpp\n"
             "QString DemoEnhancePlugin::id() const {\n"
             "    return QString::fromUtf8(\"示例增强\");  // 或 AppConfig 常量\n"
             "}\n"
             "BaseBlock *DemoEnhancePlugin::createBlock(QWidget *parent) {\n"
             "    return new DemoEnhanceBlock(parent);\n"
             "}")

    add_para(doc, "3.5.3 纯算法层（与 UI 分离）", bold=True)
    add_para(doc,
             "建议将 OpenCV 计算放在独立命名空间/函数中，不依赖 QWidget。"
             "Block 只负责读控件参数并调用算法，便于单测与复用。")
    add_code(doc,
             "// demoenhance.h/.cpp（伪代码）\n"
             "namespace DemoEnhanceAlgorithm {\n"
             "    // 输入/输出约定：BGR 三通道 Mat，与项目其它算法一致\n"
             "    cv::Mat apply(const cv::Mat &bgr, double strength);\n"
             "}")

    add_para(doc, "3.5.4 处理块（BaseBlock 子类）", bold=True)
    add_para(doc,
             "构造阶段：setupTitle → 向 contentLayout 添加参数控件 → trackParamWidget。"
             "process 阶段遵循统一五步（见 3.6）。参数序列化须先调用基类 saveParams/loadParams。")
    add_code(doc,
             "// process 关键骨架（与现有 BinarizationBlock 同构）\n"
             "QPixmap DemoEnhanceBlock::process(const QPixmap &input,\n"
             "                                 const QList<RoiInfo> &rois)\n"
             "{\n"
             "    if (input.isNull()) return input;\n"
             "    cv::Mat src = ImageConverter::pixmapToMatRGB(input);\n"
             "    cv::cvtColor(src, src, cv::COLOR_RGB2BGR);\n"
             "    const double k = m_strengthSpin->value();\n"
             "    cv::Mat result = RoiProcess::apply(src, rois, [&](const cv::Mat &m) {\n"
             "        return DemoEnhanceAlgorithm::apply(m, k);\n"
             "    });\n"
             "    return ImageConverter::matToPixmap(result);\n"
             "}\n"
             "\n"
             "QString DemoEnhanceBlock::blockName() const {\n"
             "    return QString::fromUtf8(\"示例增强\");  // 必须与 Plugin::id() 一致\n"
             "}\n"
             "\n"
             "QJsonObject DemoEnhanceBlock::saveParams() const {\n"
             "    QJsonObject o = BaseBlock::saveParams(); // 含 name / enabled\n"
             "    o.insert(\"strength\", m_strengthSpin->value());\n"
             "    return o;\n"
             "}")

    add_para(doc, "3.5.5 接入、验证与发布检查单", bold=True)
    add_table(doc, ["检查项", "通过标准"], [
        ["编译", "生成 bin/*/plugins/block_demoenhance.dll，且依赖同级 blocksdk.dll"],
        ["发现", "重启后左侧列表出现「示例增强」；或「设置→添加插件」热加载成功"],
        ["执行", "拖入链后调参，画布结果变化；耗时标签更新"],
        ["ROI", "绘制 ROI 后仅区域内变化，区外保持上一环像素"],
        ["持久化", "导出链再导入、换图再切回，块类型与参数可还原"],
        ["宿主能力（若需要）", "若实现类似 Otsu 的原图统计，经 host() 取原图，勿写死 Widget"],
    ], [4, 11.5])
    add_para(doc,
             "稳定 id 一经对外使用（会话文件、导出链）即视为协议的一部分，"
             "不应再改名；显示文案可通过 retranslateUi 国际化。")

    # ----- 3.6 -----
    add_heading(doc, "3.6 处理与数据约定", 2, *bm["s36"])
    add_para(doc,
             "约定的价值在于：所有插件遵守同一套语义后，"
             "ROI、撤销、会话、导入导出才能在主程序侧一次实现、全体受益。")

    add_para(doc, "（1）process 五步骨架", bold=True)
    add_table(doc, ["步骤", "动作", "说明"], [
        ["1", "空图防护", "input 为空则原样返回，避免 OpenCV 断言"],
        ["2", "色彩空间", "QPixmap→RGB Mat→BGR（算法层统一 BGR）"],
        ["3", "读 UI 参数", "每次 process 读最新控件值"],
        ["4", "ROI 合成", "RoiProcess::apply；空列表=全图；非空=并集蒙版局部合成"],
        ["5", "写回", "BGR Mat→QPixmap，交给下一启用块"],
    ], [1.5, 3.5, 10.5])
    add_para(doc,
             "RoiProcess::apply 语义：先对整图执行算法函数，再用 ROI 并集蒙版将结果贴回——"
             "蒙版内为处理后像素，蒙版外保持输入。禁止「只裁一块算出小图当整图返回」，"
             "否则会破坏流水线尺寸约定。")

    add_para(doc, "（2）信号与撤销", bold=True)
    add_para(doc,
             "paramsAboutToChange：值尚未改变时发出，供 Widget 压入结构快照；"
             "paramsChanged：值已更新后发出，驱动 requestReprocess。"
             "loadParams 时应对控件 blockSignals，避免恢复参数时误触发重算与误压栈。")

    add_para(doc, "（3）JSON 数据形态", bold=True)
    add_table(doc, ["层级", "关键字段", "说明"], [
        ["单块", "name, enabled, …", "name = 稳定 id；其余为块自定义参数"],
        ["单图会话", "blocks[], rois[]", "链顺序即执行顺序；ROI 可空"],
        ["落盘文件", "version, sessions{}",
         "sessions 的键为图片绝对路径；当前 version = 1"],
    ], [3, 4.5, 8])
    add_code(doc,
             "// 会话落盘结构（示意）\n"
             "{\n"
             "  \"version\": 1,\n"
             "  \"sessions\": {\n"
             "    \"D:/img/a.png\": {\n"
             "      \"blocks\": [ {\"name\":\"二值化处理\",\"enabled\":true,\"lower\":127,\"upper\":255} ],\n"
             "      \"rois\": [ /* RoiInfo JSON… */ ]\n"
             "    }\n"
             "  }\n"
             "}")

    add_para(doc, "（4）IBlockHost 使用示例（Otsu）", bold=True)
    add_para(doc,
             "二值化块在自动阈值时不使用链上中间图，而经宿主取原图与 ROI（实现摘录）：")
    add_code(doc,
             "IBlockHost *h = host();\n"
             "if (!h || !h->hostHasImage()) { /* 提示先打开图片 */ return; }\n"
             "cv::Mat src = ImageConverter::pixmapToMatRGB(h->hostOriginalImage());\n"
             "// … 按 h->hostCurrentRois() 收窄区域后计算 Otsu，写回阈值控件")

    # ----- 3.7 -----
    add_heading(doc, "3.7 工程交付与维护要点", 2, *bm["s37"])
    add_para(doc,
             "对交付与后续维护，建议把握下列基线，避免「能编译但不能运行」或「改一处牵全身」。")
    add_table(doc, ["类别", "要点"], [
        ["运行入口", "使用 bin/debug 或 bin/release 下的 opencv.exe；"
         "须旁路存在 blocksdk.dll 与 plugins/block_*.dll"],
        ["构建顺序", "blocksdk → plugins → opencv_app；改 SDK 后须全量重建插件与主程序"],
        ["配置一致", "OpenCV 路径在主程序 .pro 与 plugin_common.pri 保持一致"],
        ["协议稳定", "块 id、会话 JSON version、IBlockPlugin IID 变更需评估兼容性"],
        ["可观测性", "关键路径写 AppLogger；问题优先查阅 exe/logs/"],
        ["扩展原则", "新增能力优先落在新插件；除非改交互范式，否则避免改 ImageProcessor 核心循环"],
    ], [3.2, 12.3])
    add_table(doc, ["现有内置插件", "产物", "order"], [
        ["二值化处理", "block_binarization.dll", "10"],
        ["形态学处理", "block_morphology.dll", "20"],
        ["滤波处理", "block_filter.dll", "30"],
        ["灰度变换", "block_graytransform.dll", "40"],
        ["伪彩色处理", "block_pseudocolor.dll", "50"],
        ["灰度共生矩阵", "block_glcm.dll", "60"],
    ], [5, 6, 4])
    add_para(doc,
             "文档与实现同步方式：修改 docs/generate_spec_docx.py 后执行 "
             "python docs/generate_spec_docx.py，重新生成本稿。"
             "架构或接口变更时，应同步更新本章 3.2～3.3 与插件示例，保证上报材料与代码一致。")

    doc.add_page_break()
    add_heading(doc, "附录 术语", 1, *bm["app"])
    add_table(doc, ["术语", "含义", "说明"], [
        ["处理链", "按序排列的处理块", "顺序即流水线执行顺序"],
        ["处理块", "BaseBlock 子类实例", "含参数 UI 与 process；由插件创建"],
        ["算法插件", "实现 IBlockPlugin 的 DLL", "位于 exe/plugins/block_*.dll"],
        ["blocksdk", "插件 SDK 动态库", "提供基类、接口与 PluginManager"],
        ["多 ROI 并集", "多个感兴趣区域", "蒙版 OR；空=全图"],
        ["会话 ImageSession", "单图的链+ROI 快照", "内存哈希 + app_sessions.json"],
        ["撤销", "结构快照回退", "Ctrl+Z；深度≤40；不含 Redo"],
        ["GLCM", "灰度共生矩阵", "纹理特征分析块，图像直通"],
        ["重算 reprocess", "按当前链与 ROI 生成结果", "入口 onApplyProcessing"],
        ["热添加插件", "运行中加载单个 DLL", "设置→添加插件"],
    ], [3.2, 4.8, 7.5])

    add_para(doc,
             "（完）文档与当前源码目录 core、blocks、blocksdk、plugins、roi、utils、config 对应；"
             "若实现变更，请同步更新第一章需求表与第二章模块/接口说明，并重跑生成脚本。",
             size=10)

    doc.save(str(OUT))
    return OUT


def main():
    figures = {
        "block": make_block_diagram(),
        "flow": make_flowchart(),
        "class": make_class_diagram(),
        "inherit": make_block_inherit_diagram(),
        "plugin_dev": make_plugin_dev_diagram(),
        "dirs": make_dirs_diagram(),
    }
    out = build_doc(figures)
    print("OK:", out)


if __name__ == "__main__":
    main()
